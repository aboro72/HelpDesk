#!/usr/bin/env python3
"""
Word-Dokument Generator für Aboro-IT Handbücher
Konvertiert Markdown-Dateien zu professionell formatierten Word-Dokumenten
"""

import os
import re
from pathlib import Path
try:
    from docx import Document
    from docx.shared import Inches, Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.style import WD_STYLE_TYPE
    from docx.oxml.shared import OxmlElement, qn
    from docx.shared import RGBColor
except ImportError:
    print("❌ python-docx nicht installiert. Bitte installieren mit:")
    print("pip install python-docx")
    exit(1)

class WordDocGenerator:
    def __init__(self):
        self.base_path = Path(__file__).parent
        self.output_dir = self.base_path / "word_docs"
        self.output_dir.mkdir(exist_ok=True)
        
        # Aboro-IT Farben
        self.red_color = RGBColor(255, 68, 68)  # #FF4444
        self.black_color = RGBColor(51, 51, 51)  # #333333
        self.gray_color = RGBColor(128, 128, 128)  # #808080
    
    def create_styles(self, doc):
        """Erstellt Aboro-IT Styles für das Dokument"""
        styles = doc.styles
        
        # Heading 1 Style
        try:
            h1_style = styles['Heading 1']
        except KeyError:
            h1_style = styles.add_style('Heading 1', WD_STYLE_TYPE.PARAGRAPH)
        
        h1_font = h1_style.font
        h1_font.name = 'Calibri'
        h1_font.size = Pt(16)
        h1_font.bold = True
        h1_font.color.rgb = self.red_color
        
        # Heading 2 Style
        try:
            h2_style = styles['Heading 2']
        except KeyError:
            h2_style = styles.add_style('Heading 2', WD_STYLE_TYPE.PARAGRAPH)
        
        h2_font = h2_style.font
        h2_font.name = 'Calibri'
        h2_font.size = Pt(14)
        h2_font.bold = True
        h2_font.color.rgb = self.black_color
        
        # Heading 3 Style
        try:
            h3_style = styles['Heading 3']
        except KeyError:
            h3_style = styles.add_style('Heading 3', WD_STYLE_TYPE.PARAGRAPH)
        
        h3_font = h3_style.font
        h3_font.name = 'Calibri'
        h3_font.size = Pt(12)
        h3_font.bold = True
        h3_font.color.rgb = self.black_color
        
        # Code Style
        try:
            code_style = styles['Code']
        except KeyError:
            code_style = styles.add_style('Code', WD_STYLE_TYPE.PARAGRAPH)
        
        code_font = code_style.font
        code_font.name = 'Courier New'
        code_font.size = Pt(10)
        code_font.color.rgb = self.gray_color
    
    def add_title_page(self, doc, title, subtitle):
        """Erstellt Titelseite mit Aboro-IT Logo"""
        # Logo Platzhalter (in echter Implementierung Logo hinzufügen)
        logo_para = doc.add_paragraph()
        logo_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        logo_run = logo_para.add_run("ABORO-IT")
        logo_run.font.size = Pt(36)
        logo_run.font.bold = True
        logo_run.font.color.rgb = self.red_color
        
        # Titel
        title_para = doc.add_paragraph()
        title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_run = title_para.add_run(title)
        title_run.font.size = Pt(24)
        title_run.font.bold = True
        title_run.font.color.rgb = self.black_color
        
        # Untertitel
        subtitle_para = doc.add_paragraph()
        subtitle_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        subtitle_run = subtitle_para.add_run(subtitle)
        subtitle_run.font.size = Pt(16)
        subtitle_run.font.color.rgb = self.gray_color
        
        # Abstand
        doc.add_paragraph("\n\n")
        
        # Firmen-Info
        info_para = doc.add_paragraph()
        info_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        info_run = info_para.add_run("Professionelle IT-Lösungen für Ihr Unternehmen\nhttps://aboro-it.net")
        info_run.font.size = Pt(12)
        info_run.font.color.rgb = self.gray_color
        
        # Seite 2 beginnen
        doc.add_page_break()
    
    def add_footer(self, doc):
        """Fügt Aboro-IT Footer hinzu"""
        section = doc.sections[0]
        footer = section.footer
        footer_para = footer.paragraphs[0]
        footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        footer_run = footer_para.add_run("© 2025 Aboro-IT | Professionelle IT-Lösungen | https://aboro-it.net")
        footer_run.font.size = Pt(9)
        footer_run.font.color.rgb = self.gray_color
    
    def parse_markdown_to_word(self, markdown_content, doc):
        """Konvertiert Markdown-Inhalt zu Word-Format"""
        lines = markdown_content.split('\n')
        in_code_block = False
        code_language = ""
        
        for line in lines:
            line = line.rstrip()
            
            # Skip Logo-Placeholder und erste Überschrift (schon auf Titelseite)
            if line.startswith('![Aboro-IT Logo]') or line.startswith('# '):
                continue
            if line.startswith('## ') and 'Aboro-IT Helpdesk System' in line:
                continue
            
            # Code Blocks
            if line.startswith('```'):
                in_code_block = not in_code_block
                if in_code_block:
                    code_language = line[3:].strip()
                continue
            
            if in_code_block:
                code_para = doc.add_paragraph(line)
                code_para.style = 'Code'
                continue
            
            # Überschriften
            if line.startswith('### '):
                heading = line[4:].strip()
                para = doc.add_paragraph(heading)
                para.style = 'Heading 3'
            elif line.startswith('## '):
                heading = line[3:].strip()
                para = doc.add_paragraph(heading)
                para.style = 'Heading 2'
            elif line.startswith('# '):
                heading = line[2:].strip()
                para = doc.add_paragraph(heading)
                para.style = 'Heading 1'
            
            # Horizontale Linien (Seitenumbruch)
            elif line.strip() == '---':
                doc.add_page_break()
            
            # Listen
            elif line.startswith('- ') or line.startswith('* '):
                item_text = line[2:].strip()
                # Emojis und Formatierung beibehalten
                item_text = self.clean_markdown_formatting(item_text)
                para = doc.add_paragraph(item_text, style='List Bullet')
            
            elif re.match(r'^\d+\. ', line):
                item_text = re.sub(r'^\d+\. ', '', line).strip()
                item_text = self.clean_markdown_formatting(item_text)
                para = doc.add_paragraph(item_text, style='List Number')
            
            # Tabellen (vereinfacht)
            elif '|' in line and not line.startswith('|--'):
                # Tabellen-Behandlung (vereinfacht für Demo)
                cells = [cell.strip() for cell in line.split('|')[1:-1]]
                if cells:
                    para = doc.add_paragraph(' | '.join(cells))
                    para.style = 'Normal'
            
            # Normaler Text
            elif line.strip():
                if not line.startswith('|--'):  # Skip Tabellen-Separator
                    text = self.clean_markdown_formatting(line)
                    if text.strip():
                        para = doc.add_paragraph(text)
                        para.style = 'Normal'
            
            # Leerzeile
            else:
                doc.add_paragraph('')
    
    def clean_markdown_formatting(self, text):
        """Entfernt/konvertiert Markdown-Formatierung"""
        # Bold **text** -> text (Word-Formatierung später)
        text = re.sub(r'\*\*(.*?)\*\*', r'\1', text)
        # Italic *text* -> text
        text = re.sub(r'\*(.*?)\*', r'\1', text)
        # Code `text` -> text
        text = re.sub(r'`(.*?)`', r'\1', text)
        # Links [text](url) -> text
        text = re.sub(r'\[([^\]]+)\]\([^\)]+\)', r'\1', text)
        
        return text
    
    def create_word_document(self, markdown_file, output_name, title, subtitle):
        """Erstellt Word-Dokument aus Markdown-Datei"""
        print(f"Erstelle {output_name}...")
        
        # Markdown-Datei lesen
        try:
            with open(self.base_path / markdown_file, 'r', encoding='utf-8') as f:
                content = f.read()
        except FileNotFoundError:
            print(f"FEHLER: Datei {markdown_file} nicht gefunden!")
            return False
        except Exception as e:
            print(f"FEHLER beim Lesen von {markdown_file}: {e}")
            return False
        
        # Word-Dokument erstellen
        doc = Document()
        
        # Styles erstellen
        self.create_styles(doc)
        
        # Titelseite hinzufügen
        self.add_title_page(doc, title, subtitle)
        
        # Footer hinzufügen
        self.add_footer(doc)
        
        # Markdown-Inhalt konvertieren
        self.parse_markdown_to_word(content, doc)
        
        # Dokument speichern
        output_path = self.output_dir / output_name
        try:
            doc.save(output_path)
            print(f"ERFOLG: {output_name} erstellt: {output_path}")
            return True
        except Exception as e:
            print(f"FEHLER beim Speichern von {output_name}: {e}")
            return False
    
    def create_all_documents(self):
        """Erstellt alle Aboro-IT Handbücher als Word-Dokumente"""
        documents = [
            {
                'markdown_file': 'ADMINISTRATOR_HANDBUCH.md',
                'output_name': 'Aboro-IT_Administrator_Handbuch.docx',
                'title': 'Administrator Handbuch',
                'subtitle': 'Aboro-IT Helpdesk System - Vollstaendige Systemverwaltung'
            },
            {
                'markdown_file': 'BENUTZERHANDBUCH.md',
                'output_name': 'Aboro-IT_Support_Agent_Handbuch.docx',
                'title': 'Support Agent Handbuch',
                'subtitle': 'Aboro-IT Helpdesk System - Vollstaendige Agent-Anleitung'
            },
            {
                'markdown_file': 'ENTWICKLER_HANDBUCH.md',
                'output_name': 'Aboro-IT_Entwickler_Handbuch.docx',
                'title': 'Entwickler Handbuch',
                'subtitle': 'Aboro-IT Helpdesk System - Technische Dokumentation'
            },
            {
                'markdown_file': 'DOKUMENTATION_INDEX.md',
                'output_name': 'Aboro-IT_Dokumentation_Index.docx',
                'title': 'Dokumentation Index',
                'subtitle': 'Aboro-IT Helpdesk System - Vollstaendige Dokumentationsuebersicht'
            }
        ]
        
        print("Starte Word-Dokument-Erstellung fuer Aboro-IT...")
        print(f"Output-Verzeichnis: {self.output_dir}")
        print("=" * 60)
        
        success_count = 0
        for doc_info in documents:
            if self.create_word_document(**doc_info):
                success_count += 1
        
        print("=" * 60)
        print(f"ERFOLG: {success_count}/{len(documents)} Dokumente erfolgreich erstellt!")
        
        if success_count == len(documents):
            print("Alle Aboro-IT Handbuecher wurden erfolgreich als Word-Dateien erstellt!")
            print(f"Dateien verfuegbar in: {self.output_dir}")
        else:
            print("WARNUNG: Einige Dokumente konnten nicht erstellt werden. Pruefen Sie die Fehlermeldungen oben.")
        
        return success_count == len(documents)

def main():
    """Hauptfunktion"""
    print("Aboro-IT Word-Dokument Generator")
    print("=" * 50)
    
    # Prüfen ob python-docx installiert ist
    try:
        import docx
    except ImportError:
        print("FEHLER: python-docx ist nicht installiert!")
        print("Bitte installieren Sie es mit: pip install python-docx")
        return False
    
    # Generator erstellen und ausführen
    generator = WordDocGenerator()
    return generator.create_all_documents()

if __name__ == "__main__":
    success = main()
    input("\nDrücken Sie Enter um zu beenden...")
    exit(0 if success else 1)