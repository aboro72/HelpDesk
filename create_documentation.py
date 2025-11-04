from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from datetime import datetime

def create_user_manual():
    """Create user manual document"""
    doc = Document()

    # Set up margins
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    # Title
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title.add_run('ML Gruppe Helpdesk System')
    title_run.font.size = Pt(24)
    title_run.font.bold = True
    title_run.font.color.rgb = RGBColor(0, 51, 102)

    # Subtitle
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle_run = subtitle.add_run('Benutzerhandbuch')
    subtitle_run.font.size = Pt(18)
    subtitle_run.font.color.rgb = RGBColor(51, 102, 153)

    # Date
    date_para = doc.add_paragraph()
    date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    date_run = date_para.add_run(f'Version 1.0 - {datetime.now().strftime("%d.%m.%Y")}')
    date_run.font.size = Pt(10)
    date_run.font.italic = True

    doc.add_paragraph()  # Spacing

    # Table of Contents
    toc_heading = doc.add_heading('Inhaltsverzeichnis', level=1)
    toc_items = [
        '1. Einführung',
        '2. Erste Schritte',
        '3. Tickets erstellen und verwalten',
        '4. Ticketsuche und Filter',
        '5. Ticketkommentare',
        '6. Knowledge Base nutzen',
        '7. Häufig gestellte Fragen'
    ]
    for item in toc_items:
        doc.add_paragraph(item, style='List Bullet')

    doc.add_page_break()

    # 1. Introduction
    doc.add_heading('1. Einführung', level=1)
    doc.add_paragraph(
        'Das ML Gruppe Helpdesk System ist eine webbasierte Anwendung zur Verwaltung von Support-Tickets. '
        'Mit diesem System können Sie Support-Anfragen erstellen, verfolgen und lösen.'
    )

    doc.add_heading('Funktionsübersicht', level=2)
    features = [
        'Ticketmanagement: Erstellen, bearbeiten und schließen Sie Tickets',
        'Ticketsuche: Finden Sie schnell Ihre Tickets mit Filteroptionen',
        'Knowledge Base: Zugriff auf Lösungsartikel und häufig gestellte Fragen',
        'Statistiken: Übersicht über Support-Anfragen und Trends',
        'Benachrichtigungen: Automatische Email-Benachrichtigungen über Ticketänderungen'
    ]
    for feature in features:
        doc.add_paragraph(feature, style='List Bullet')

    # 2. Getting Started
    doc.add_heading('2. Erste Schritte', level=1)

    doc.add_heading('Anmeldung', level=2)
    doc.add_paragraph(
        'Öffnen Sie Ihren Webbrowser und navigieren Sie zur Helpdesk-URL. '
        'Geben Sie Ihre Email-Adresse und Ihr Passwort ein. Nach dem ersten Login müssen Sie Ihr Passwort ändern.'
    )

    doc.add_heading('Dashboard', level=2)
    doc.add_paragraph(
        'Nach der Anmeldung sehen Sie das Hauptdashboard mit einer Übersicht Ihrer Tickets. '
        'Sie können von hier aus neue Tickets erstellen oder bestehende Tickets verwalten.'
    )

    doc.add_heading('Navigation', level=2)
    nav_items = [
        'Tickets: Liste aller Ihrer Tickets',
        'Ticket erstellen: Neues Ticket einreichen',
        'Knowledge Base: Artikel und Lösungen durchsuchen',
        'Profil: Ihre Kontoeinstellungen'
    ]
    for item in nav_items:
        doc.add_paragraph(item, style='List Bullet')

    # 3. Tickets erstellen und verwalten
    doc.add_heading('3. Tickets erstellen und verwalten', level=1)

    doc.add_heading('Ein neues Ticket erstellen', level=2)
    doc.add_paragraph('Klicken Sie auf "Ticket erstellen" und füllen Sie das Formular aus:')
    ticket_fields = [
        'Titel: Kurze Zusammenfassung des Problems',
        'Beschreibung: Detaillierte Erklärung des Problems',
        'Kategorie: Wählen Sie die passende Kategorie',
        'Priorität: Setzen Sie die Dringlichkeit fest',
        'Mobiler Klassenraum (optional): Falls zutreffend'
    ]
    for field in ticket_fields:
        doc.add_paragraph(field, style='List Bullet')

    doc.add_heading('Prioritätsstufen', level=2)
    priorities = [
        'Low: Nicht dringend',
        'Medium: Normale Priorität',
        'High: Dringend',
        'Critical: Sehr dringend'
    ]
    for priority in priorities:
        doc.add_paragraph(priority, style='List Bullet')

    doc.add_heading('Ticketstatus', level=2)
    statuses = [
        'Open: Ticket wurde gerade erstellt',
        'In Progress: Support-Team arbeitet am Ticket',
        'Resolved: Problem wurde gelöst',
        'Closed: Ticket ist abgeschlossen'
    ]
    for status in statuses:
        doc.add_paragraph(status, style='List Bullet')

    # 4. Search and Filter
    doc.add_heading('4. Ticketsuche und Filter', level=1)

    doc.add_heading('Tickets durchsuchen', level=2)
    doc.add_paragraph(
        'Auf der Ticketliste können Sie Ihre Tickets nach Status, Priorität oder Kategorie filtern. '
        'Verwenden Sie die Suchfunktion, um schnell ein bestimmtes Ticket zu finden.'
    )

    doc.add_heading('Filter verwenden', level=2)
    filters = [
        'Status: Filtern Sie nach Ticketstatus',
        'Priorität: Zeigen Sie nur bestimmte Prioritäten',
        'Kategorie: Filtern Sie nach Problemtyp'
    ]
    for f in filters:
        doc.add_paragraph(f, style='List Bullet')

    # 5. Comments
    doc.add_heading('5. Ticketkommentare', level=1)
    doc.add_paragraph(
        'Sie können Kommentare zu Ihren Tickets hinzufügen, um zusätzliche Informationen bereitzustellen '
        'oder auf Fragen des Support-Teams zu antworten.'
    )

    doc.add_heading('Kommentar hinzufügen', level=2)
    comment_steps = [
        'Öffnen Sie das Ticket',
        'Scrollen Sie zum Kommentarbereich',
        'Geben Sie Ihren Kommentar ein',
        'Klicken Sie auf "Kommentar hinzufügen"'
    ]
    for step in comment_steps:
        doc.add_paragraph(step, style='List Bullet')

    # 6. Knowledge Base
    doc.add_heading('6. Knowledge Base nutzen', level=1)
    doc.add_paragraph(
        'Die Knowledge Base enthält hilfreiche Artikel und Lösungen für häufige Probleme. '
        'Sie können Artikel durchsuchen oder nach Kategorie browsen.'
    )

    doc.add_heading('Artikel durchsuchen', level=2)
    kb_features = [
        'Verwenden Sie die Suchfunktion, um Artikel zu finden',
        'Filtern Sie nach Kategorie',
        'Lesen Sie verwandte Artikel'
    ]
    for feature in kb_features:
        doc.add_paragraph(feature, style='List Bullet')

    # 7. FAQ
    doc.add_heading('7. Häufig gestellte Fragen', level=1)

    doc.add_heading('Wie lange dauert es, bis mein Ticket bearbeitet wird?', level=2)
    doc.add_paragraph(
        'Die Bearbeitungszeit hängt von der Priorität ab. '
        'Critical-Tickets werden innerhalb von 1 Stunde bearbeitet, High-Tickets innerhalb von 4 Stunden.'
    )

    doc.add_heading('Kann ich mein Passwort ändern?', level=2)
    doc.add_paragraph('Ja, gehen Sie zu Ihrem Profil und klicken Sie auf "Passwort ändern".')

    doc.add_heading('Wie lösche ich ein Ticket?', level=2)
    doc.add_paragraph('Geschlossene Tickets können nicht gelöscht werden. Kontaktieren Sie den Administrator.')

    doc.add_heading('Warum erhalte ich keine Email-Benachrichtigungen?', level=2)
    doc.add_paragraph(
        'Überprüfen Sie Ihre Email-Adresse in den Kontoeinstellungen. '
        'Benachrichtigungen werden möglicherweise in Ihrem Spam-Ordner gefiltert.'
    )

    # Footer
    doc.add_paragraph()
    footer_para = doc.add_paragraph()
    footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer_run = footer_para.add_run('© 2025 ML Gruppe Helpdesk System - Alle Rechte vorbehalten')
    footer_run.font.size = Pt(8)
    footer_run.font.italic = True

    doc.save('Benutzerhandbuch_ML_Helpdesk.docx')
    return 'Benutzerhandbuch_ML_Helpdesk.docx'

def create_admin_manual():
    """Create administration manual document"""
    doc = Document()

    # Set up margins
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    # Title
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title.add_run('ML Gruppe Helpdesk System')
    title_run.font.size = Pt(24)
    title_run.font.bold = True
    title_run.font.color.rgb = RGBColor(0, 51, 102)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle_run = subtitle.add_run('Administrationshandbuch')
    subtitle_run.font.size = Pt(18)
    subtitle_run.font.color.rgb = RGBColor(51, 102, 153)

    date_para = doc.add_paragraph()
    date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    date_run = date_para.add_run(f'Version 1.0 - {datetime.now().strftime("%d.%m.%Y")}')
    date_run.font.size = Pt(10)
    date_run.font.italic = True

    doc.add_paragraph()

    # Table of Contents
    doc.add_heading('Inhaltsverzeichnis', level=1)
    toc_items = [
        '1. Systemübersicht',
        '2. Benutzerverwaltung',
        '3. Ticketkonfiguration',
        '4. Knowledge Base Administration',
        '5. Kategorien und Status',
        '6. Support-Level und Team-Lead',
        '7. Statistiken und Berichte',
        '8. Sicherheit und Datenschutz',
        '9. Backup und Wartung',
        '10. Troubleshooting'
    ]
    for item in toc_items:
        doc.add_paragraph(item, style='List Bullet')

    doc.add_page_break()

    # 1. System Overview
    doc.add_heading('1. Systemübersicht', level=1)
    doc.add_paragraph(
        'Das ML Gruppe Helpdesk System ist eine Django-basierte Anwendung zur Verwaltung von Support-Tickets. '
        'Das System unterstützt mehrere Benutzerrollen mit unterschiedlichen Berechtigungen.'
    )

    doc.add_heading('Benutzerrollen', level=2)
    roles = [
        'Administrator: Vollständiger Zugriff auf alle Funktionen',
        'Support Agent: Kann Tickets verwalten und Knowledge Base Articles erstellen',
        'Team Lead (Level 4): Kann Tickets an andere Agents zuweisen',
        'Customer: Kann nur ihre eigenen Tickets sehen'
    ]
    for role in roles:
        doc.add_paragraph(role, style='List Bullet')

    # 2. User Management
    doc.add_heading('2. Benutzerverwaltung', level=1)

    doc.add_heading('Neue Benutzer erstellen', level=2)
    doc.add_paragraph(
        'Nur Administratoren können neue Benutzer erstellen. Gehen Sie zum Admin-Panel '
        'und navigieren Sie zu "Benutzer" > "Neuer Benutzer".'
    )

    user_fields = [
        'Email: Eindeutige Email-Adresse',
        'Username: Eindeutiger Benutzername',
        'Vorname und Nachname: Vollständiger Name',
        'Rolle: Wählen Sie die entsprechende Rolle',
        'Support Level: Nur für Support Agents (1-4)',
        'Telefonnummer: Kontaktinformation'
    ]
    for field in user_fields:
        doc.add_paragraph(field, style='List Bullet')

    doc.add_heading('Passwortverwaltung', level=2)
    doc.add_paragraph(
        'Neue Kunden erhalten ein initiales Passwort: P@ssw0rd123. '
        'Sie müssen das Passwort bei ihrer ersten Anmeldung ändern.'
    )

    doc.add_heading('Benutzer deaktivieren', level=2)
    doc.add_paragraph(
        'Benutzer können deaktiviert (nicht gelöscht) werden, indem Sie das Feld "Aktiv" deaktivieren. '
        'Deaktivierte Benutzer können sich nicht mehr anmelden.'
    )

    # 3. Ticket Configuration
    doc.add_heading('3. Ticketkonfiguration', level=1)

    doc.add_heading('Ticketkategorien', level=2)
    doc.add_paragraph(
        'Kategorien helfen bei der Klassifizierung von Tickets. Sie können im Admin-Panel '
        'neue Kategorien hinzufügen oder bestehende bearbeiten.'
    )

    doc.add_heading('SLA (Service Level Agreement)', level=2)
    doc.add_paragraph(
        'Das System unterstützt automatische SLA-Verwaltung basierend auf Priorität:'
    )
    sla_items = [
        'Critical: 1 Stunde Reaktionszeit',
        'High: 4 Stunden Reaktionszeit',
        'Medium: 8 Stunden Reaktionszeit',
        'Low: 24 Stunden Reaktionszeit'
    ]
    for item in sla_items:
        doc.add_paragraph(item, style='List Bullet')

    doc.add_heading('Mobile Classroom', level=2)
    doc.add_paragraph(
        'Mobile Classrooms können im Admin-Panel verwaltet werden. '
        'Jeder Classroom hat einen Standort und kann mehrere Tickets haben.'
    )

    # 4. Knowledge Base
    doc.add_heading('4. Knowledge Base Administration', level=1)

    doc.add_heading('Artikel erstellen', level=2)
    doc.add_paragraph(
        'Support Agents und Administratoren können Knowledge Base Artikel erstellen. '
        'Artikel sollten häufige Probleme und deren Lösungen dokumentieren.'
    )

    kb_fields = [
        'Titel: Aussagekräftiger Titel',
        'Kategorie: Thematische Zuordnung',
        'Inhalt: Detaillierte Anleitung (mit RichText-Editor)',
        'Tags: Für bessere Auffindbarkeit',
        'Veröffentlicht: Status des Artikels'
    ]
    for field in kb_fields:
        doc.add_paragraph(field, style='List Bullet')

    doc.add_heading('Artikel-Verwaltung', level=2)
    doc.add_paragraph(
        'Ändern Sie im Admin-Panel die Kategorie oder den Status von Artikeln. '
        'Archivierte Artikel werden nicht mehr in der Suche angezeigt.'
    )

    # 5. Categories and Status
    doc.add_heading('5. Kategorien und Status', level=1)

    doc.add_heading('Ticketstatus verwalten', level=2)
    doc.add_paragraph('Standardstatus sind:')
    statuses = [
        'Open: Neu erstelltes Ticket',
        'In Progress: Unter Bearbeitung',
        'Resolved: Problem gelöst',
        'Closed: Ticket abgeschlossen'
    ]
    for status in statuses:
        doc.add_paragraph(status, style='List Bullet')

    # 6. Support Levels
    doc.add_heading('6. Support-Level und Team-Lead', level=1)

    doc.add_heading('Support Levels', level=2)
    doc.add_paragraph('Es gibt 4 Support Levels für Agents:')
    levels = [
        'Level 1: Junior Support (kann nur einfache Tickets lösen)',
        'Level 2: Senior Support (normale Tickets)',
        'Level 3: Expert Support (komplexe Probleme)',
        'Level 4: Team Lead (kann Tickets zuweisen und delegieren)'
    ]
    for level in levels:
        doc.add_paragraph(level, style='List Bullet')

    doc.add_heading('Team Lead Funktionen', level=2)
    doc.add_paragraph('Team Leads mit Level 4 können:')
    tl_features = [
        'Tickets an andere Agents zuweisen',
        'Tickets eskalieren',
        'Team-Statistiken anschauen',
        'Andere Agents überwachen'
    ]
    for feature in tl_features:
        doc.add_paragraph(feature, style='List Bullet')

    # 7. Statistics
    doc.add_heading('7. Statistiken und Berichte', level=1)
    doc.add_paragraph(
        'Das System bietet umfangreiche Statistiken zur Überwachung der Support-Qualität. '
        'Nur Administratoren und Team Leads haben Zugriff auf Statistiken.'
    )

    doc.add_heading('Verfügbare Statistiken', level=2)
    stats = [
        'Tickets pro Trainer/Kunde',
        'Hochprioräts-Tickets',
        'Durchschnittliche Auflösungszeit',
        'Häufigste Probleme/Kategorien',
        'Mobile Classroom Probleme',
        'Prioritätsverteilung'
    ]
    for stat in stats:
        doc.add_paragraph(stat, style='List Bullet')

    # 8. Security
    doc.add_heading('8. Sicherheit und Datenschutz', level=1)

    doc.add_heading('Passwortrichtlinien', level=2)
    doc.add_paragraph(
        'Alle Passwörter werden verschlüsselt gespeichert. '
        'Benutzer sollten regelmäßig ihre Passwörter ändern.'
    )

    doc.add_heading('Datenzugriff', level=2)
    doc.add_paragraph(
        'Kunden können nur ihre eigenen Tickets sehen. '
        'Support Agents und Administratoren sehen alle Tickets.'
    )

    doc.add_heading('Email-Benachrichtigungen', level=2)
    doc.add_paragraph(
        'Das System sendet automatische Benachrichtigungen. '
        'Überprüfen Sie regelmäßig Email-Konfiguration und SMTP-Einstellungen.'
    )

    # 9. Backup and Maintenance
    doc.add_heading('9. Backup und Wartung', level=1)

    doc.add_heading('Regelmäßige Backups', level=2)
    doc.add_paragraph(
        'Erstellen Sie täglich Backups der Datenbank. '
        'Speichern Sie Backups an einem sicheren Ort außerhalb des Systems.'
    )

    doc.add_heading('Systemwartung', level=2)
    doc.add_paragraph(
        'Führen Sie regelmäßig folgende Wartungsaufgaben durch:'
    )
    maintenance = [
        'Datenbankoptimierung',
        'Log-Datei-Archivierung',
        'Benutzer-Account-Überprüfung',
        'Security-Updates durchführen'
    ]
    for task in maintenance:
        doc.add_paragraph(task, style='List Bullet')

    # 10. Troubleshooting
    doc.add_heading('10. Troubleshooting', level=1)

    doc.add_heading('Häufige Probleme', level=2)

    doc.add_heading('Tickets werden nicht angezeigt', level=3)
    doc.add_paragraph(
        'Überprüfen Sie, ob der Benutzer die richtige Rolle hat. '
        'Administratoren sehen alle Tickets, Kunden nur ihre eigenen.'
    )

    doc.add_heading('Email-Benachrichtigungen funktionieren nicht', level=3)
    doc.add_paragraph(
        'Überprüfen Sie die SMTP-Konfiguration in settings.py. '
        'Stellen Sie sicher, dass der Email-Server erreichbar ist.'
    )

    doc.add_heading('Datenbankfehler', level=3)
    doc.add_paragraph(
        'Führen Sie Migrationen durch: python manage.py migrate. '
        'Überprüfen Sie die Datenbankverbindung.'
    )

    # Footer
    doc.add_paragraph()
    footer_para = doc.add_paragraph()
    footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer_run = footer_para.add_run('© 2025 ML Gruppe Helpdesk System - Administrationshandbuch')
    footer_run.font.size = Pt(8)
    footer_run.font.italic = True

    doc.save('Administrationshandbuch_ML_Helpdesk.docx')
    return 'Administrationshandbuch_ML_Helpdesk.docx'

def create_dev_manual():
    """Create developer manual document"""
    doc = Document()

    # Set up margins
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    # Title
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title.add_run('ML Gruppe Helpdesk System')
    title_run.font.size = Pt(24)
    title_run.font.bold = True
    title_run.font.color.rgb = RGBColor(0, 51, 102)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle_run = subtitle.add_run('Entwicklerhandbuch')
    subtitle_run.font.size = Pt(18)
    subtitle_run.font.color.rgb = RGBColor(51, 102, 153)

    date_para = doc.add_paragraph()
    date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    date_run = date_para.add_run(f'Version 1.0 - {datetime.now().strftime("%d.%m.%Y")}')
    date_run.font.size = Pt(10)
    date_run.font.italic = True

    doc.add_paragraph()

    # Table of Contents
    doc.add_heading('Inhaltsverzeichnis', level=1)
    toc_items = [
        '1. Projektstruktur',
        '2. Technologie Stack',
        '3. Installation und Setup',
        '4. Datenmodelle',
        '5. Views und URL-Routing',
        '6. Formularverarbeitung',
        '7. AJAX API-Endpoints',
        '8. Authentifizierung und Autorisierung',
        '9. Signale und Event-Handling',
        '10. Testing und Debugging',
        '11. Deployment'
    ]
    for item in toc_items:
        doc.add_paragraph(item, style='List Bullet')

    doc.add_page_break()

    # 1. Project Structure
    doc.add_heading('1. Projektstruktur', level=1)
    doc.add_paragraph('Das Projekt folgt einer typischen Django-Struktur:')

    structure = [
        'mini-helpdesk/ - Hauptprojektverzeichnis',
        '  ├── helpdesk/ - Django Projekteinstellungen (settings.py, urls.py, wsgi.py)',
        '  ├── apps/ - Django-Anwendungen',
        '  │   ├── accounts/ - Benutzerverwaltung',
        '  │   ├── tickets/ - Ticketmanagement',
        '  │   ├── knowledge/ - Knowledge Base',
        '  │   └── main/ - Dashboard und Übersichten',
        '  ├── templates/ - HTML-Templates',
        '  ├── static/ - CSS, JavaScript, Bilder',
        '  ├── manage.py - Django Management Skript',
        '  └── requirements.txt - Abhängigkeiten'
    ]
    for item in structure:
        if item.startswith('  '):
            doc.add_paragraph(item)
        else:
            doc.add_paragraph(item, style='List Bullet')

    # 2. Technology Stack
    doc.add_heading('2. Technologie Stack', level=1)

    doc.add_heading('Backend', level=2)
    backend = [
        'Django 5.0+ - Web Framework',
        'Python 3.11+ - Programmiersprache',
        'SQLite/PostgreSQL - Datenbank',
        'Celery - Asynchrone Task Queue (optional)',
        'Gunicorn - WSGI HTTP Server'
    ]
    for tech in backend:
        doc.add_paragraph(tech, style='List Bullet')

    doc.add_heading('Frontend', level=2)
    frontend = [
        'Bootstrap 5 - CSS Framework',
        'JavaScript (Vanilla JS) - Client-side Logik',
        'TinyMCE - Rich Text Editor',
        'HTML5 - Markup'
    ]
    for tech in frontend:
        doc.add_paragraph(tech, style='List Bullet')

    doc.add_heading('Abhängigkeiten (wichtigste)', level=2)
    deps = [
        'django - Web Framework',
        'django-crispy-forms - Form Styling',
        'pillow - Bildverarbeitung',
        'python-docx - Word Document Generation',
        'requests - HTTP Library'
    ]
    for dep in deps:
        doc.add_paragraph(dep, style='List Bullet')

    # 3. Installation
    doc.add_heading('3. Installation und Setup', level=1)

    doc.add_heading('Vorbedingungen', level=2)
    doc.add_paragraph('Python 3.11+')
    doc.add_paragraph('pip (Python Package Manager)')
    doc.add_paragraph('Git (für Versionskontrolle)')

    doc.add_heading('Setup-Schritte', level=2)
    steps = [
        '1. Repository klonen: git clone <repo-url>',
        '2. In Verzeichnis navigieren: cd mini-helpdesk',
        '3. Virtual Environment erstellen: python -m venv .venv',
        '4. Virtual Environment aktivieren: .venv\\Scripts\\activate (Windows)',
        '5. Abhängigkeiten installieren: pip install -r requirements.txt',
        '6. Migrationen durchführen: python manage.py migrate',
        '7. Superuser erstellen: python manage.py createsuperuser',
        '8. Server starten: python manage.py runserver'
    ]
    for step in steps:
        doc.add_paragraph(step, style='List Number')

    # 4. Data Models
    doc.add_heading('4. Datenmodelle', level=1)

    doc.add_heading('User Model (apps/accounts/models.py)', level=2)
    doc.add_paragraph('Erweitertes Django User Model mit zusätzlichen Feldern:')
    user_fields = [
        'email - Email-Adresse',
        'phone - Telefonnummer',
        'role - Benutzerrolle (customer, support_agent, admin)',
        'support_level - Support Level (1-4 für Agents)',
        'force_password_change - Passwort-Wechsel erzwingen'
    ]
    for field in user_fields:
        doc.add_paragraph(field, style='List Bullet')

    doc.add_heading('Ticket Model (apps/tickets/models.py)', level=2)
    doc.add_paragraph('Speichert Support-Tickets mit:')
    ticket_fields = [
        'ticket_number - Eindeutige Ticketnummer',
        'title - Tickeltitel',
        'description - Detaillierte Beschreibung',
        'status - Ticketstatus (open, in_progress, resolved, closed)',
        'priority - Priorität (low, medium, high, critical)',
        'category - Problemkategorie',
        'created_by - Kunde der das Ticket erstellt hat',
        'assigned_to - Support Agent der das Ticket bearbeitet',
        'mobile_classroom - Verknüpfter mobiler Klassenraum',
        'created_at / updated_at - Zeitstempel',
        'sla_due_at - Service Level Agreement Deadline'
    ]
    for field in ticket_fields:
        doc.add_paragraph(field, style='List Bullet')

    doc.add_heading('TicketComment Model', level=2)
    doc.add_paragraph('Speichert Kommentare zu Tickets:')
    comment_fields = [
        'ticket - Verknüpfung zum Ticket',
        'author - Autor des Kommentars',
        'content - Kommentarinhalt',
        'is_internal - Nur sichtbar für Support-Team',
        'created_at - Zeitstempel'
    ]
    for field in comment_fields:
        doc.add_paragraph(field, style='List Bullet')

    doc.add_heading('KnowledgeArticle Model', level=2)
    doc.add_paragraph('Knowledge Base Artikel:')
    kb_fields = [
        'title - Artikeltitel',
        'slug - URL-freundlicher Identifier',
        'content - HTML-Inhalt (TinyMCE Editor)',
        'category - Themenkategorie',
        'is_published - Veröffentlichungsstatus',
        'created_by - Autor',
        'created_at - Erstellungsdatum'
    ]
    for field in kb_fields:
        doc.add_paragraph(field, style='List Bullet')

    # 5. Views
    doc.add_heading('5. Views und URL-Routing', level=1)

    doc.add_heading('Ticket Views (apps/tickets/views.py)', level=2)
    views = [
        'ticket_list() - Listet Tickets basierend auf Benutzerrolle',
        'ticket_create() - Erstellt neues Ticket (mit Agent-Support)',
        'ticket_detail() - Zeigt Ticketdetails und Kommentare',
        'ticket_assign() - Weist Ticket Agent zu',
        'ticket_escalate() - Eskaliert Ticket mit Email-Benachrichtigung',
        'ticket_close() - Schließt Ticket und sendet Geschichte',
        'statistics_dashboard() - Zeigt Statistiken für Agents/Admins',
        'search_customers_api() - AJAX API für Kundensuche'
    ]
    for view in views:
        doc.add_paragraph(view, style='List Bullet')

    # 6. Forms
    doc.add_heading('6. Formularverarbeitung', level=1)

    doc.add_heading('AgentTicketCreateForm', level=2)
    doc.add_paragraph(
        'Formular für Support Agents zum Erstellen von Tickets im Namen von Kunden. '
        'Merkmale:'
    )
    form_features = [
        'customer_search - AJAX-gestützte Kundensuche',
        'customer_email - Email des Kunden (automatisch gefüllt)',
        'customer_first_name - Vorname (automatisch gefüllt)',
        'customer_last_name - Nachname (automatisch gefüllt)',
        'customer_phone - Telefonnummer (automatisch gefüllt)',
        'ticket_fields - Titel, Beschreibung, Kategorie, Priorität'
    ]
    for feature in form_features:
        doc.add_paragraph(feature, style='List Bullet')

    # 7. AJAX
    doc.add_heading('7. AJAX API-Endpoints', level=1)

    doc.add_heading('Kundensuche API', level=2)
    doc.add_paragraph('Endpoint: /tickets/api/search-customers/')
    doc.add_paragraph('Methode: GET')
    doc.add_paragraph('Parameter: q (Suchstring)')
    doc.add_paragraph('Rückgabe: JSON mit Kundenliste')

    doc.add_heading('Response Format', level=2)
    response = '''
{
  "results": [
    {
      "id": 123,
      "name": "Hans Hirschfeld",
      "first_name": "Hans",
      "last_name": "Hirschfeld",
      "email": "hans.hirschfeld@example.com",
      "phone": "+49 40 555-1234",
      "display": "Hans Hirschfeld (hans.hirschfeld@example.com)"
    }
  ]
}
    '''
    doc.add_paragraph(response)

    # 8. Authentication
    doc.add_heading('8. Authentifizierung und Autorisierung', level=1)

    doc.add_heading('Decorator: @login_required', level=2)
    doc.add_paragraph('Sichert Views, so dass nur angemeldete Benutzer Zugriff haben.')

    doc.add_heading('Permission Check: can_access_ticket()', level=2)
    doc.add_paragraph(
        'Benutzerdefinierte Methode im User Model. '
        'Customers sehen nur ihre eigenen Tickets, Agents/Admins sehen alle.'
    )

    doc.add_heading('Role-based Access Control', level=2)
    access = [
        'customer - Kann nur eigene Tickets sehen',
        'support_agent - Kann alle Tickets sehen und bearbeiten',
        'admin - Vollständiger Zugriff auf alles'
    ]
    for role in access:
        doc.add_paragraph(role, style='List Bullet')

    # 9. Signals
    doc.add_heading('9. Signale und Event-Handling', level=1)

    doc.add_heading('post_save Signal', level=2)
    doc.add_paragraph(
        'Automatische Benachrichtigungen wenn Tickets erstellt werden. '
        'Sendet Emails an alle Agents.'
    )

    doc.add_heading('Escalation Handler', level=2)
    doc.add_paragraph(
        'Automatische Benachrichtigung mit Dringlichkeitsangabe '
        'wenn Tickets eskaliert werden.'
    )

    # 10. Testing
    doc.add_heading('10. Testing und Debugging', level=1)

    doc.add_heading('Unit Tests', level=2)
    doc.add_paragraph('Tests können in apps/*/tests.py geschrieben werden.')

    doc.add_heading('Alle Tests ausführen', level=2)
    doc.add_paragraph('python manage.py test')

    doc.add_heading('Django Shell', level=2)
    doc.add_paragraph('python manage.py shell')
    doc.add_paragraph('Ermöglicht Interaktion mit Datenbank und Models.')

    doc.add_heading('Debugging', level=2)
    debug_tips = [
        'Nutzen Sie print() oder logging.info() für Debug-Ausgaben',
        'Überprüfen Sie Django Logs in der Konsole',
        'Nutzen Sie Django Debug Toolbar für Performance-Analyse',
        'Überprüfen Sie Browser Console für JavaScript-Fehler'
    ]
    for tip in debug_tips:
        doc.add_paragraph(tip, style='List Bullet')

    # 11. Deployment
    doc.add_heading('11. Deployment', level=1)

    doc.add_heading('Production Setup', level=2)
    doc.add_paragraph(
        'Für Produktionsumgebung sind folgende Schritte empfohlen:'
    )

    deploy_steps = [
        '1. Nutzen Sie einen Production-WSGI-Server (Gunicorn)',
        '2. Setzen Sie DEBUG=False in settings.py',
        '3. Verwenden Sie eine externe Datenbank (PostgreSQL)',
        '4. Konfigurieren Sie ALLOWED_HOSTS korrekt',
        '5. Nutzen Sie HTTPS mit SSL-Zertifikat',
        '6. Setzen Sie Umgebungsvariablen für sensitive Daten',
        '7. Konfigurieren Sie Email-Server für Benachrichtigungen',
        '8. Implementieren Sie regelmäßige Backups'
    ]
    for step in deploy_steps:
        doc.add_paragraph(step, style='List Number')

    doc.add_heading('Environment Variables', level=2)
    vars = [
        'SECRET_KEY - Django Secret Key',
        'DEBUG - False für Production',
        'ALLOWED_HOSTS - Erlaubte Hostnames',
        'DATABASE_URL - Datenbankverbindungsstring',
        'EMAIL_HOST - SMTP Server',
        'EMAIL_PORT - SMTP Port (meist 587)',
        'EMAIL_USE_TLS - True für TLS',
        'EMAIL_HOST_USER - Email-Adresse',
        'EMAIL_HOST_PASSWORD - Email-Passwort'
    ]
    for var in vars:
        doc.add_paragraph(var, style='List Bullet')

    # Footer
    doc.add_paragraph()
    footer_para = doc.add_paragraph()
    footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer_run = footer_para.add_run('© 2025 ML Gruppe Helpdesk System - Entwicklerhandbuch')
    footer_run.font.size = Pt(8)
    footer_run.font.italic = True

    doc.save('Entwicklerhandbuch_ML_Helpdesk.docx')
    return 'Entwicklerhandbuch_ML_Helpdesk.docx'

if __name__ == '__main__':
    print("Erstelle Dokumentationen...\n")

    file1 = create_user_manual()
    print(f"[OK] {file1} erstellt")

    file2 = create_admin_manual()
    print(f"[OK] {file2} erstellt")

    file3 = create_dev_manual()
    print(f"[OK] {file3} erstellt")

    print(f"\n[SUCCESS] Alle 3 Dokumentationen erfolgreich erstellt!")
    print(f"\nDateien:")
    print(f"  - Benutzerhandbuch_ML_Helpdesk.docx")
    print(f"  - Administrationshandbuch_ML_Helpdesk.docx")
    print(f"  - Entwicklerhandbuch_ML_Helpdesk.docx")
