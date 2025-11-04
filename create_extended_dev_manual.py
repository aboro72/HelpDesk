from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from datetime import datetime

def create_extended_dev_manual():
    """Create extended developer manual with database structure and app development"""
    doc = Document()

    # Set up margins
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(0.8)
        section.bottom_margin = Inches(0.8)
        section.left_margin = Inches(0.8)
        section.right_margin = Inches(0.8)

    # Title
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title.add_run('ML Gruppe Helpdesk System')
    title_run.font.size = Pt(24)
    title_run.font.bold = True
    title_run.font.color.rgb = RGBColor(0, 51, 102)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle_run = subtitle.add_run('Entwicklerhandbuch - Erweiterte Edition')
    subtitle_run.font.size = Pt(18)
    subtitle_run.font.color.rgb = RGBColor(51, 102, 153)

    date_para = doc.add_paragraph()
    date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    date_run = date_para.add_run(f'Version 2.0 - {datetime.now().strftime("%d.%m.%Y")}')
    date_run.font.size = Pt(10)
    date_run.font.italic = True

    doc.add_paragraph()

    # Table of Contents
    doc.add_heading('Inhaltsverzeichnis', level=1)
    toc_items = [
        '1. Projektstruktur',
        '2. Technologie Stack',
        '3. Installation und Setup',
        '4. Datenbankstruktur (Detailliert)',
        '5. Models - Tiefgehende Analyse',
        '6. Views und URL-Routing',
        '7. Formularverarbeitung',
        '8. AJAX API-Endpoints',
        '9. Authentifizierung und Autorisierung',
        '10. Signale und Event-Handling',
        '11. Neue Django Apps erstellen',
        '12. Testing und Debugging',
        '13. Performance Optimierung',
        '14. Deployment',
        '15. Best Practices'
    ]
    for item in toc_items:
        doc.add_paragraph(item, style='List Bullet')

    doc.add_page_break()

    # 1. Project Structure
    doc.add_heading('1. Projektstruktur', level=1)

    structure = '''
mini-helpdesk/
├── manage.py                      # Django Management Skript
├── requirements.txt               # Python Abhängigkeiten
├── .gitignore                     # Git Ignore Datei
├── README.md                      # Projekt Dokumentation
│
├── helpdesk/                      # Django Projekteinstellungen
│   ├── __init__.py
│   ├── settings.py                # Haupt-Konfigurationsdatei
│   ├── urls.py                    # Haupt-URL Router
│   ├── asgi.py                    # ASGI Interface
│   ├── wsgi.py                    # WSGI Interface
│   └── celery.py                  # Celery Konfiguration
│
├── apps/                          # Django Anwendungen
│   ├── accounts/                  # Benutzerverwaltung
│   │   ├── migrations/
│   │   ├── __init__.py
│   │   ├── admin.py
│   │   ├── apps.py
│   │   ├── forms.py
│   │   ├── models.py
│   │   ├── tests.py
│   │   ├── urls.py
│   │   └── views.py
│   │
│   ├── tickets/                   # Ticketmanagement
│   │   ├── migrations/
│   │   ├── management/commands/   # Custom Django Commands
│   │   ├── __init__.py
│   │   ├── admin.py
│   │   ├── apps.py
│   │   ├── forms.py
│   │   ├── models.py
│   │   ├── signals.py
│   │   ├── tests.py
│   │   ├── urls.py
│   │   ├── views.py
│   │   ├── ai_service.py          # Claude AI Integration
│   │   └── [other utilities]
│   │
│   ├── knowledge/                 # Knowledge Base
│   │   ├── migrations/
│   │   ├── __init__.py
│   │   ├── admin.py
│   │   ├── apps.py
│   │   ├── forms.py
│   │   ├── models.py
│   │   ├── tests.py
│   │   ├── urls.py
│   │   └── views.py
│   │
│   └── main/                      # Dashboard & Übersicht
│       ├── migrations/
│       ├── __init__.py
│       ├── admin.py
│       ├── apps.py
│       ├── tests.py
│       ├── urls.py
│       └── views.py
│
├── templates/                     # HTML Templates
│   ├── base.html                  # Basis Template
│   ├── accounts/
│   ├── tickets/
│   ├── knowledge/
│   └── main/
│
├── static/                        # CSS, JavaScript, Bilder
│   ├── css/
│   ├── js/
│   ├── images/
│   └── vendor/
│
├── media/                         # Benutzerdateien (Uploads)
│
└── docs/                          # Dokumentation
    ├── Benutzerhandbuch_ML_Helpdesk.docx
    ├── Administrationshandbuch_ML_Helpdesk.docx
    └── Entwicklerhandbuch_ML_Helpdesk.docx
    '''
    doc.add_paragraph(structure)

    # 2. Technology Stack
    doc.add_heading('2. Technologie Stack', level=1)

    doc.add_heading('Backend Framework', level=2)
    doc.add_paragraph('Django 5.0+: Web Framework mit ORM, Admin, Authentication')
    doc.add_paragraph('Python 3.11+: Programmiersprache')

    doc.add_heading('Datenbanken', level=2)
    doc.add_paragraph('SQLite: Entwicklung (Standard)')
    doc.add_paragraph('PostgreSQL: Produktion (empfohlen)')
    doc.add_paragraph('MySQL/MariaDB: Alternative für Produktion')

    doc.add_heading('Caching & Message Queue', level=2)
    doc.add_paragraph('Redis: Cache und Session Store')
    doc.add_paragraph('Celery: Asynchrone Task Queue')
    doc.add_paragraph('RabbitMQ: Message Broker (optional)')

    doc.add_heading('Frontend', level=2)
    doc.add_paragraph('Bootstrap 5: CSS Framework')
    doc.add_paragraph('HTML5: Markup')
    doc.add_paragraph('JavaScript (Vanilla): Client-side Logik')
    doc.add_paragraph('TinyMCE: Rich Text Editor')

    doc.add_heading('Python Libraries', level=2)
    libs = [
        'django-crispy-forms: Form Rendering',
        'pillow: Bildverarbeitung',
        'python-docx: Word Document Generation',
        'requests: HTTP Requests',
        'python-decouple: Umgebungsvariablen',
        'gunicorn: WSGI Server',
        'psycopg2: PostgreSQL Adapter',
        'mysqlclient: MySQL Adapter',
        'django-redis: Redis Cache Backend',
        'celery: Task Queue',
        'python-dateutil: Datum/Zeit Utilities'
    ]
    for lib in libs:
        doc.add_paragraph(lib, style='List Bullet')

    # 3. Installation
    doc.add_heading('3. Installation und Setup', level=1)

    doc.add_heading('Entwicklungsumgebung einrichten', level=2)
    setup_steps = '''
# 1. Repository klonen
git clone <repo-url>
cd mini-helpdesk

# 2. Virtual Environment erstellen
python3 -m venv .venv

# 3. Aktivieren
# Windows:
.venv\\Scripts\\activate
# Linux/Mac:
source .venv/bin/activate

# 4. Dependencies installieren
pip install -r requirements.txt

# 5. Environment Variablen setzen
cp .env.example .env
# Bearbeiten Sie .env mit Ihren Werten

# 6. Datenbank initialisieren
python manage.py migrate

# 7. Superuser erstellen
python manage.py createsuperuser

# 8. Testdaten laden (optional)
python manage.py loaddata fixtures/initial_data.json

# 9. Static Files sammeln
python manage.py collectstatic --noinput

# 10. Development Server starten
python manage.py runserver
    '''
    doc.add_paragraph(setup_steps)

    doc.add_page_break()

    # 4. Database Structure
    doc.add_heading('4. Datenbankstruktur (Detailliert)', level=1)

    doc.add_heading('Entity-Relationship Diagramm (Logisch)', level=2)
    erd = '''
User (auth_user + Erweiterung)
├── email (UNIQUE)
├── password (gehashed)
├── first_name
├── last_name
├── phone
├── role (customer, support_agent, admin)
├── support_level (1-4, nur für Agents)
├── is_active
├── created_at
└── force_password_change

Category (tickets_category)
├── name (UNIQUE)
├── description
├── is_active
└── created_at

Ticket (tickets_ticket)
├── ticket_number (UNIQUE)
├── title
├── description
├── status (open, in_progress, resolved, closed)
├── priority (low, medium, high, critical)
├── created_by (FK -> User)
├── assigned_to (FK -> User, nullable)
├── category (FK -> Category)
├── mobile_classroom (FK -> MobileClassroom, nullable)
├── created_at
├── updated_at
├── resolved_at (nullable)
├── closed_at (nullable)
└── sla_due_at

TicketComment (tickets_ticketcomment)
├── ticket (FK -> Ticket)
├── author (FK -> User)
├── content
├── is_internal (boolean)
├── created_at
└── updated_at

MobileClassroom (tickets_mobileclassroom)
├── name
├── location (FK -> MobileClassroomLocation)
├── is_active
└── created_at

MobileClassroomLocation (tickets_mobileclassroomlocation)
├── name
├── address
└── is_active

KnowledgeArticle (knowledge_knowledgearticle)
├── title
├── slug (UNIQUE)
├── content (HTML)
├── category (FK -> KnowledgeCategory)
├── is_published (boolean)
├── created_by (FK -> User)
├── created_at
├── updated_at
└── view_count

KnowledgeCategory (knowledge_knowledgecategory)
├── name
├── description
└── is_active
    '''
    doc.add_paragraph(erd)

    doc.add_heading('Datenbankschema - SQL Struktur', level=2)
    schema = '''
-- Benutzer (erweitert)
CREATE TABLE auth_user (
    id INTEGER PRIMARY KEY,
    username VARCHAR(150) UNIQUE NOT NULL,
    email VARCHAR(254) UNIQUE NOT NULL,
    password VARCHAR(128) NOT NULL,
    first_name VARCHAR(150),
    last_name VARCHAR(150),
    phone VARCHAR(20),
    role VARCHAR(20) DEFAULT 'customer',
    support_level INTEGER DEFAULT 1,
    is_active BOOLEAN DEFAULT TRUE,
    force_password_change BOOLEAN DEFAULT FALSE,
    created_at TIMESTAMP DEFAULT NOW(),
    updated_at TIMESTAMP DEFAULT NOW()
);

-- Kategorien
CREATE TABLE tickets_category (
    id INTEGER PRIMARY KEY,
    name VARCHAR(100) UNIQUE NOT NULL,
    description TEXT,
    is_active BOOLEAN DEFAULT TRUE,
    created_at TIMESTAMP DEFAULT NOW()
);

-- Tickets
CREATE TABLE tickets_ticket (
    id INTEGER PRIMARY KEY,
    ticket_number VARCHAR(20) UNIQUE NOT NULL,
    title VARCHAR(200) NOT NULL,
    description TEXT NOT NULL,
    status VARCHAR(20) DEFAULT 'open',
    priority VARCHAR(20) DEFAULT 'medium',
    created_by_id INTEGER NOT NULL REFERENCES auth_user(id),
    assigned_to_id INTEGER REFERENCES auth_user(id),
    category_id INTEGER REFERENCES tickets_category(id),
    mobile_classroom_id INTEGER REFERENCES tickets_mobileclassroom(id),
    created_at TIMESTAMP DEFAULT NOW(),
    updated_at TIMESTAMP DEFAULT NOW(),
    resolved_at TIMESTAMP,
    closed_at TIMESTAMP,
    sla_due_at TIMESTAMP
);

-- Ticket Kommentare
CREATE TABLE tickets_ticketcomment (
    id INTEGER PRIMARY KEY,
    ticket_id INTEGER NOT NULL REFERENCES tickets_ticket(id),
    author_id INTEGER NOT NULL REFERENCES auth_user(id),
    content TEXT NOT NULL,
    is_internal BOOLEAN DEFAULT FALSE,
    created_at TIMESTAMP DEFAULT NOW(),
    updated_at TIMESTAMP DEFAULT NOW()
);

-- Knowledge Base
CREATE TABLE knowledge_knowledgearticle (
    id INTEGER PRIMARY KEY,
    title VARCHAR(200) NOT NULL,
    slug VARCHAR(200) UNIQUE NOT NULL,
    content TEXT NOT NULL,
    category_id INTEGER REFERENCES knowledge_knowledgecategory(id),
    is_published BOOLEAN DEFAULT FALSE,
    created_by_id INTEGER NOT NULL REFERENCES auth_user(id),
    created_at TIMESTAMP DEFAULT NOW(),
    updated_at TIMESTAMP DEFAULT NOW(),
    view_count INTEGER DEFAULT 0
);

-- Indices für Performance
CREATE INDEX idx_ticket_created_by ON tickets_ticket(created_by_id);
CREATE INDEX idx_ticket_assigned_to ON tickets_ticket(assigned_to_id);
CREATE INDEX idx_ticket_status ON tickets_ticket(status);
CREATE INDEX idx_ticket_priority ON tickets_ticket(priority);
CREATE INDEX idx_ticket_created_at ON tickets_ticket(created_at);
CREATE INDEX idx_comment_ticket ON tickets_ticketcomment(ticket_id);
CREATE INDEX idx_comment_author ON tickets_ticketcomment(author_id);
CREATE INDEX idx_article_published ON knowledge_knowledgearticle(is_published);
    '''
    doc.add_paragraph(schema)

    doc.add_heading('Beziehungen zwischen Tabellen', level=2)
    doc.add_paragraph('1:N Beziehungen:')
    doc.add_paragraph('User -> Ticket (created_by): Ein User kann viele Tickets erstellen', style='List Bullet')
    doc.add_paragraph('User -> Ticket (assigned_to): Ein Agent kann viele Tickets bearbeiten', style='List Bullet')
    doc.add_paragraph('Ticket -> TicketComment: Ein Ticket kann viele Kommentare haben', style='List Bullet')
    doc.add_paragraph('Category -> Ticket: Eine Kategorie kann viele Tickets haben', style='List Bullet')
    doc.add_paragraph('User -> KnowledgeArticle: Ein User kann viele Artikel schreiben', style='List Bullet')

    # 5. Models
    doc.add_heading('5. Models - Tiefgehende Analyse', level=1)

    doc.add_heading('User Model (apps/accounts/models.py)', level=2)
    user_model = '''
from django.contrib.auth.models import AbstractUser

class User(AbstractUser):
    ROLE_CHOICES = [
        ('customer', 'Customer'),
        ('support_agent', 'Support Agent'),
        ('admin', 'Administrator'),
    ]

    SUPPORT_LEVEL_CHOICES = [
        (1, 'Level 1 - Junior'),
        (2, 'Level 2 - Senior'),
        (3, 'Level 3 - Expert'),
        (4, 'Level 4 - Team Lead'),
    ]

    # Zusätzliche Felder
    phone = models.CharField(max_length=20, blank=True)
    role = models.CharField(max_length=20, choices=ROLE_CHOICES, default='customer')
    support_level = models.IntegerField(choices=SUPPORT_LEVEL_CHOICES, default=1)
    force_password_change = models.BooleanField(default=False)
    created_at = models.DateTimeField(auto_now_add=True)
    updated_at = models.DateTimeField(auto_now=True)

    class Meta:
        db_table = 'auth_user'
        ordering = ['created_at']

    def __str__(self):
        return f"{self.get_full_name()} ({self.get_role_display()})"

    @property
    def full_name(self):
        return f"{self.first_name} {self.last_name}".strip()

    def can_access_ticket(self, ticket):
        """Check if user can access this ticket"""
        if self.role == 'admin':
            return True
        if self.role == 'support_agent':
            return True
        # Customers can only see their own tickets
        return ticket.created_by == self
    '''
    doc.add_paragraph(user_model)

    doc.add_heading('Ticket Model (apps/tickets/models.py)', level=2)
    ticket_model = '''
class Ticket(models.Model):
    STATUS_CHOICES = [
        ('open', 'Open'),
        ('in_progress', 'In Progress'),
        ('resolved', 'Resolved'),
        ('closed', 'Closed'),
    ]

    PRIORITY_CHOICES = [
        ('low', 'Low'),
        ('medium', 'Medium'),
        ('high', 'High'),
        ('critical', 'Critical'),
    ]

    # Wichtige Felder
    ticket_number = models.CharField(max_length=20, unique=True)
    title = models.CharField(max_length=200)
    description = models.TextField()
    status = models.CharField(max_length=20, choices=STATUS_CHOICES, default='open')
    priority = models.CharField(max_length=20, choices=PRIORITY_CHOICES, default='medium')

    # Beziehungen
    created_by = models.ForeignKey(User, on_delete=models.PROTECT, related_name='created_tickets')
    assigned_to = models.ForeignKey(User, on_delete=models.SET_NULL, null=True, blank=True,
                                    related_name='assigned_tickets')
    category = models.ForeignKey(Category, on_delete=models.SET_NULL, null=True, blank=True)
    mobile_classroom = models.ForeignKey(MobileClassroom, on_delete=models.SET_NULL,
                                         null=True, blank=True)

    # Zeitstempel
    created_at = models.DateTimeField(auto_now_add=True)
    updated_at = models.DateTimeField(auto_now=True)
    resolved_at = models.DateTimeField(null=True, blank=True)
    closed_at = models.DateTimeField(null=True, blank=True)
    sla_due_at = models.DateTimeField(null=True, blank=True)

    class Meta:
        db_table = 'tickets_ticket'
        ordering = ['-created_at']
        indexes = [
            models.Index(fields=['status']),
            models.Index(fields=['priority']),
            models.Index(fields=['created_at']),
        ]

    def __str__(self):
        return f"{self.ticket_number} - {self.title}"

    def set_priority_based_sla(self):
        """Set SLA due date based on priority"""
        from datetime import timedelta
        sla_hours = {'low': 24, 'medium': 8, 'high': 4, 'critical': 1}
        self.sla_due_at = self.created_at + timedelta(hours=sla_hours[self.priority])

    def get_history_as_text(self):
        """Get ticket history for email"""
        history = f"Ticket {self.ticket_number}: {self.title}\\n"
        history += f"Status: {self.get_status_display()}\\n"
        # ... weitere Details ...
        return history
    '''
    doc.add_paragraph(ticket_model)

    doc.add_heading('TicketComment Model', level=2)
    comment_model = '''
class TicketComment(models.Model):
    ticket = models.ForeignKey(Ticket, on_delete=models.CASCADE, related_name='comments')
    author = models.ForeignKey(User, on_delete=models.PROTECT)
    content = models.TextField()
    is_internal = models.BooleanField(default=False)
    created_at = models.DateTimeField(auto_now_add=True)
    updated_at = models.DateTimeField(auto_now=True)

    class Meta:
        db_table = 'tickets_ticketcomment'
        ordering = ['created_at']

    def __str__(self):
        return f"Comment on {self.ticket.ticket_number} by {self.author}"
    '''
    doc.add_paragraph(comment_model)

    doc.add_heading('KnowledgeArticle Model', level=2)
    kb_model = '''
class KnowledgeArticle(models.Model):
    title = models.CharField(max_length=200)
    slug = models.SlugField(unique=True)
    content = models.TextField()  # HTML Content from TinyMCE
    category = models.ForeignKey(KnowledgeCategory, on_delete=models.SET_NULL,
                                 null=True, blank=True)
    is_published = models.BooleanField(default=False)
    created_by = models.ForeignKey(User, on_delete=models.PROTECT)
    created_at = models.DateTimeField(auto_now_add=True)
    updated_at = models.DateTimeField(auto_now=True)
    view_count = models.IntegerField(default=0)

    class Meta:
        db_table = 'knowledge_knowledgearticle'
        ordering = ['-created_at']

    def __str__(self):
        return self.title

    def save(self, *args, **kwargs):
        """Auto-generate slug from title"""
        if not self.slug:
            from django.utils.text import slugify
            self.slug = slugify(self.title)
        super().save(*args, **kwargs)

    def increment_views(self):
        """Increment view counter"""
        self.view_count += 1
        self.save(update_fields=['view_count'])
    '''
    doc.add_paragraph(kb_model)

    doc.add_page_break()

    # 6. Views
    doc.add_heading('6. Views und URL-Routing', level=1)

    doc.add_heading('URL Routing Struktur (urls.py)', level=2)
    urls = '''
# helpdesk/urls.py (Haupt URL Router)
from django.contrib import admin
from django.urls import path, include
from django.conf import settings
from django.conf.urls.static import static

urlpatterns = [
    path('admin/', admin.site.urls),
    path('', include('apps.main.urls', namespace='main')),
    path('accounts/', include('apps.accounts.urls', namespace='accounts')),
    path('tickets/', include('apps.tickets.urls', namespace='tickets')),
    path('kb/', include('apps.knowledge.urls', namespace='knowledge')),
]

if settings.DEBUG:
    urlpatterns += static(settings.MEDIA_URL, document_root=settings.MEDIA_ROOT)
    urlpatterns += static(settings.STATIC_URL, document_root=settings.STATIC_ROOT)

# apps/tickets/urls.py (Ticket App URLs)
from django.urls import path
from . import views

app_name = 'tickets'

urlpatterns = [
    path('', views.ticket_list, name='list'),
    path('create/', views.ticket_create, name='create'),
    path('<int:pk>/', views.ticket_detail, name='detail'),
    path('<int:pk>/assign/', views.ticket_assign, name='assign'),
    path('<int:pk>/escalate/', views.ticket_escalate, name='escalate'),
    path('<int:pk>/close/', views.ticket_close, name='close'),
    path('api/search-customers/', views.search_customers_api, name='search_customers_api'),
]
    '''
    doc.add_paragraph(urls)

    doc.add_heading('View Funktion Beispiel', level=2)
    view_example = '''
from django.shortcuts import render, redirect, get_object_or_404
from django.contrib.auth.decorators import login_required
from django.contrib import messages
from .models import Ticket, TicketComment
from .forms import TicketCreateForm

@login_required
def ticket_detail(request, pk):
    """View ticket details and add comments"""
    # Security Check: Ensure user has access
    ticket = get_object_or_404(Ticket, pk=pk)
    if not request.user.can_access_ticket(ticket):
        return HttpResponseForbidden()

    # Handle comment submission
    if request.method == 'POST':
        form = TicketCommentForm(request.POST)
        if form.is_valid():
            comment = form.save(commit=False)
            comment.ticket = ticket
            comment.author = request.user
            comment.save()
            messages.success(request, 'Comment added!')
            return redirect('tickets:detail', pk=ticket.pk)
    else:
        form = TicketCommentForm()

    # Get visible comments based on user role
    if request.user.role == 'customer':
        comments = ticket.comments.filter(is_internal=False)
    else:
        comments = ticket.comments.all()

    context = {
        'ticket': ticket,
        'comments': comments,
        'form': form,
    }
    return render(request, 'tickets/detail.html', context)
    '''
    doc.add_paragraph(view_example)

    doc.add_page_break()

    # 7. Forms
    doc.add_heading('7. Formularverarbeitung', level=1)

    doc.add_heading('Form Klassen (forms.py)', level=2)
    forms = '''
from django import forms
from django.contrib.auth.forms import UserCreationForm
from .models import Ticket, TicketComment, User

class UserRegistrationForm(UserCreationForm):
    """Form für Benutzer Registrierung"""
    email = forms.EmailField(required=True)
    first_name = forms.CharField(max_length=100, required=True)
    last_name = forms.CharField(max_length=100, required=True)

    class Meta:
        model = User
        fields = ('email', 'first_name', 'last_name', 'password1', 'password2')

class TicketCreateForm(forms.ModelForm):
    """Form für Ticket Erstellung durch Kunden"""
    class Meta:
        model = Ticket
        fields = ['title', 'description', 'category', 'priority', 'mobile_classroom']
        widgets = {
            'title': forms.TextInput(attrs={'class': 'form-control'}),
            'description': forms.Textarea(attrs={'class': 'form-control', 'rows': 6}),
            'category': forms.Select(attrs={'class': 'form-control'}),
            'priority': forms.Select(attrs={'class': 'form-control'}),
            'mobile_classroom': forms.Select(attrs={'class': 'form-control'}),
        }

class TicketCommentForm(forms.ModelForm):
    """Form für Ticket Kommentare"""
    is_internal = forms.BooleanField(
        required=False,
        label='Internal comment (only visible to staff)',
        widget=forms.CheckboxInput(attrs={'class': 'form-check-input'})
    )

    class Meta:
        model = TicketComment
        fields = ['content', 'is_internal']
        widgets = {
            'content': forms.Textarea(attrs={'class': 'form-control', 'rows': 4}),
        }

    def clean(self):
        """Custom validation"""
        cleaned_data = super().clean()
        content = cleaned_data.get('content')

        if content and len(content.strip()) < 10:
            raise forms.ValidationError("Comment must be at least 10 characters long")

        return cleaned_data
    '''
    doc.add_paragraph(forms)

    doc.add_page_break()

    # 11. Creating New Apps
    doc.add_heading('11. Neue Django Apps erstellen', level=1)

    doc.add_heading('Schritt-für-Schritt Anleitung', level=2)

    doc.add_heading('1. App Gerüst erstellen', level=3)
    doc.add_paragraph('python manage.py startapp my_new_app')

    doc.add_heading('2. App in INSTALLED_APPS registrieren', level=3)
    install_apps = '''
# helpdesk/settings.py
INSTALLED_APPS = [
    'django.contrib.admin',
    'django.contrib.auth',
    # ...
    'apps.accounts',
    'apps.tickets',
    'apps.knowledge',
    'apps.main',
    'apps.my_new_app',  # Neue App hinzufügen
]
    '''
    doc.add_paragraph(install_apps)

    doc.add_heading('3. Models definieren', level=3)
    new_app_models = '''
# apps/my_new_app/models.py
from django.db import models
from apps.accounts.models import User

class MyModel(models.Model):
    """Beispiel Model für neue App"""
    name = models.CharField(max_length=100)
    description = models.TextField()
    owner = models.ForeignKey(User, on_delete=models.CASCADE)
    created_at = models.DateTimeField(auto_now_add=True)
    updated_at = models.DateTimeField(auto_now=True)

    class Meta:
        db_table = 'my_new_app_mymodel'
        ordering = ['-created_at']

    def __str__(self):
        return self.name
    '''
    doc.add_paragraph(new_app_models)

    doc.add_heading('4. Admin registrieren', level=3)
    new_app_admin = '''
# apps/my_new_app/admin.py
from django.contrib import admin
from .models import MyModel

@admin.register(MyModel)
class MyModelAdmin(admin.ModelAdmin):
    list_display = ('name', 'owner', 'created_at')
    list_filter = ('created_at', 'owner')
    search_fields = ('name', 'description')
    date_hierarchy = 'created_at'

    def get_queryset(self, request):
        """Filter basierend auf Benutzerrolle"""
        qs = super().get_queryset(request)
        if not request.user.is_superuser:
            qs = qs.filter(owner=request.user)
        return qs
    '''
    doc.add_paragraph(new_app_admin)

    doc.add_heading('5. Forms erstellen', level=3)
    new_app_forms = '''
# apps/my_new_app/forms.py
from django import forms
from .models import MyModel

class MyModelForm(forms.ModelForm):
    class Meta:
        model = MyModel
        fields = ['name', 'description']
        widgets = {
            'name': forms.TextInput(attrs={'class': 'form-control'}),
            'description': forms.Textarea(attrs={'class': 'form-control', 'rows': 4}),
        }
    '''
    doc.add_paragraph(new_app_forms)

    doc.add_heading('6. Views erstellen', level=3)
    new_app_views = '''
# apps/my_new_app/views.py
from django.shortcuts import render, redirect, get_object_or_404
from django.contrib.auth.decorators import login_required
from django.contrib import messages
from .models import MyModel
from .forms import MyModelForm

@login_required
def my_model_list(request):
    """List all MyModel objects"""
    if request.user.role == 'admin':
        objects = MyModel.objects.all()
    else:
        objects = MyModel.objects.filter(owner=request.user)

    return render(request, 'my_new_app/list.html', {'objects': objects})

@login_required
def my_model_create(request):
    """Create new MyModel"""
    if request.method == 'POST':
        form = MyModelForm(request.POST)
        if form.is_valid():
            obj = form.save(commit=False)
            obj.owner = request.user
            obj.save()
            messages.success(request, 'Created successfully!')
            return redirect('my_new_app:list')
    else:
        form = MyModelForm()

    return render(request, 'my_new_app/form.html', {'form': form})
    '''
    doc.add_paragraph(new_app_views)

    doc.add_heading('7. URLs konfigurieren', level=3)
    new_app_urls = '''
# apps/my_new_app/urls.py
from django.urls import path
from . import views

app_name = 'my_new_app'

urlpatterns = [
    path('', views.my_model_list, name='list'),
    path('create/', views.my_model_create, name='create'),
]

# In helpdesk/urls.py hinzufügen:
urlpatterns = [
    # ...
    path('my-app/', include('apps.my_new_app.urls', namespace='my_new_app')),
]
    '''
    doc.add_paragraph(new_app_urls)

    doc.add_heading('8. Migrations erstellen', level=3)
    doc.add_paragraph('python manage.py makemigrations my_new_app')
    doc.add_paragraph('python manage.py migrate my_new_app')

    doc.add_heading('9. Templates erstellen', level=3)
    doc.add_paragraph('Erstelle Template Ordner: templates/my_new_app/')
    doc.add_paragraph('Erstelle Dateien: list.html, form.html, detail.html')

    doc.add_page_break()

    # 12. Testing
    doc.add_heading('12. Testing und Debugging', level=1)

    doc.add_heading('Unit Tests schreiben', level=2)
    tests = '''
# apps/tickets/tests.py
from django.test import TestCase, Client
from django.contrib.auth import get_user_model
from .models import Ticket, Category

User = get_user_model()

class TicketModelTest(TestCase):
    def setUp(self):
        """Setup test data"""
        self.user = User.objects.create_user(
            email='test@example.com',
            password='testpass123'
        )
        self.category = Category.objects.create(name='Test Category')

    def test_ticket_creation(self):
        """Test creating a ticket"""
        ticket = Ticket.objects.create(
            ticket_number='TST-001',
            title='Test Ticket',
            description='Test description',
            created_by=self.user,
            category=self.category
        )
        self.assertTrue(Ticket.objects.filter(ticket_number='TST-001').exists())

    def test_ticket_sla_calculation(self):
        """Test SLA due date calculation"""
        ticket = Ticket.objects.create(
            ticket_number='TST-002',
            title='Test',
            description='Test',
            created_by=self.user,
            priority='high'
        )
        ticket.set_priority_based_sla()
        self.assertIsNotNone(ticket.sla_due_at)

class TicketViewTest(TestCase):
    def setUp(self):
        self.client = Client()
        self.user = User.objects.create_user(
            email='test@example.com',
            password='testpass123'
        )

    def test_ticket_list_requires_login(self):
        """Test that ticket list requires authentication"""
        response = self.client.get('/tickets/')
        self.assertEqual(response.status_code, 302)  # Redirect to login

    def test_ticket_list_authenticated(self):
        """Test ticket list for authenticated user"""
        self.client.login(email='test@example.com', password='testpass123')
        response = self.client.get('/tickets/')
        self.assertEqual(response.status_code, 200)
    '''
    doc.add_paragraph(tests)

    doc.add_heading('Tests ausführen', level=2)
    doc.add_paragraph('# Alle Tests:')
    doc.add_paragraph('python manage.py test')
    doc.add_paragraph('# Spezifische App:')
    doc.add_paragraph('python manage.py test apps.tickets')
    doc.add_paragraph('# Mit Verbosity:')
    doc.add_paragraph('python manage.py test -v 2')

    doc.add_page_break()

    # 13. Performance
    doc.add_heading('13. Performance Optimierung', level=1)

    doc.add_heading('Database Queries optimieren', level=2)
    perf_tips = '''
# SCHLECHT - N+1 Problem:
tickets = Ticket.objects.all()
for ticket in tickets:
    print(ticket.created_by.full_name)  # Extra Query für jeden Ticket!

# GUT - Select Related:
tickets = Ticket.objects.select_related('created_by', 'category')
for ticket in tickets:
    print(ticket.created_by.full_name)  # Keine extra Queries

# GUT - Prefetch Related für M2M:
comments = TicketComment.objects.prefetch_related('ticket')
    '''
    doc.add_paragraph(perf_tips)

    doc.add_heading('Caching implementieren', level=2)
    caching = '''
from django.views.decorators.cache import cache_page
from django.core.cache import cache

# View-Level Caching (60 Sekunden)
@cache_page(60)
def cached_list(request):
    items = MyModel.objects.all()
    return render(request, 'list.html', {'items': items})

# Manuelles Caching:
def get_stats():
    cache_key = 'stats_cache'
    stats = cache.get(cache_key)

    if stats is None:
        # Calculate stats
        stats = expensive_calculation()
        cache.set(cache_key, stats, 3600)  # Cache for 1 hour

    return stats
    '''
    doc.add_paragraph(caching)

    doc.add_heading('Indexierung', level=2)
    doc.add_paragraph('Füge Indizes zu häufig abgefragten Feldern hinzu:')
    doc.add_paragraph('- ticket.status')
    doc.add_paragraph('- ticket.priority')
    doc.add_paragraph('- ticket.created_at')
    doc.add_paragraph('- ticket.created_by_id')

    # 14. Deployment
    doc.add_heading('14. Deployment', level=1)

    doc.add_heading('Production Checklist', level=2)
    checklist = [
        'DEBUG = False in settings.py',
        'SECRET_KEY aus Umgebungsvariable',
        'ALLOWED_HOSTS korrekt konfiguriert',
        'Datenbank: PostgreSQL oder MySQL',
        'Static Files sammeln: collectstatic',
        'HTTPS/SSL aktiviert',
        'Backups konfiguriert',
        'Logging konfiguriert',
        'Email Server konfiguriert',
        'Redis/Cache konfiguriert',
        'Celery Worker läuft',
        'Monitoring eingerichtet'
    ]
    for item in checklist:
        doc.add_paragraph(item, style='List Bullet')

    # 15. Best Practices
    doc.add_heading('15. Best Practices', level=1)

    doc.add_heading('Code Style und Struktur', level=2)
    practices = '''
1. Verwende Black für Code Formatting:
   pip install black
   black apps/

2. Verwende flake8 für Linting:
   pip install flake8
   flake8 apps/

3. Django Code Style Guide befolgen:
   - PEP 8 für Python
   - Deskriptive Variablennamen
   - Docstrings für Funktionen
   - Type Hints wo möglich (Python 3.6+)

4. Git Workflow:
   - feature/* Branches für neue Features
   - bugfix/* Branches für Bugfixes
   - Aussagekräftige Commit Messages
   - Pull Requests für Code Review

5. Dokumentation:
   - Docstrings in Funktionen
   - Inline Comments für komplexe Logik
   - README.md aktuell halten
   - API Dokumentation (z.B. mit Swagger)

6. Security:
   - Input Validation
   - SQL Injection Prevention (ORM nutzt dies automatisch)
   - XSS Protection
   - CSRF Protection (Django eingebaut)
   - HTTPS nur
   - Secrets nicht in Code speichern
    '''
    doc.add_paragraph(practices)

    # Footer
    doc.add_paragraph()
    footer_para = doc.add_paragraph()
    footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer_run = footer_para.add_run('© 2025 ML Gruppe Helpdesk System - Erweitertes Entwicklerhandbuch')
    footer_run.font.size = Pt(8)
    footer_run.font.italic = True

    doc.save('Entwicklerhandbuch_ML_Helpdesk_ERWEITERT.docx')
    return 'Entwicklerhandbuch_ML_Helpdesk_ERWEITERT.docx'

if __name__ == '__main__':
    print("[ERSTELLE] Erweitertes Entwicklerhandbuch...\n")
    file = create_extended_dev_manual()
    print(f"[OK] {file} erstellt")
    print(f"\n[SUCCESS] Erweitertes Entwicklerhandbuch fertig!")
    print(f"\nDatei:")
    print(f"  - Entwicklerhandbuch_ML_Helpdesk_ERWEITERT.docx")
