from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from datetime import datetime

def create_extended_admin_manual():
    """Create extended administration manual with deployment options"""
    doc = Document()

    # Set up margins
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    # Title
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title.add_run('ML Gruppe Helpdesk System')
    title_run.font.size = Pt(24)
    title_run.font.bold = True
    title_run.font.color.rgb = RGBColor(0, 51, 102)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle_run = subtitle.add_run('Administrationshandbuch - Erweiterte Edition')
    subtitle_run.font.size = Pt(18)
    subtitle_run.font.color.rgb = RGBColor(51, 102, 153)

    date_para = doc.add_paragraph()
    date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    date_run = date_para.add_run(f'Version 2.0 - {datetime.now().strftime("%d.%m.%Y")}')
    date_run.font.size = Pt(10)
    date_run.font.italic = True

    doc.add_paragraph()

    # Table of Contents
    doc.add_heading('Inhaltsverzeichnis', level=1)
    toc_items = [
        '1. Systemübersicht',
        '2. Benutzerverwaltung',
        '3. Datenbank-Konfiguration',
        '4. Caching und Queue-System',
        '5. Standalone Deployment (Apache/Nginx)',
        '6. ISPConfig 3 mit Apache2',
        '7. ISPConfig 3 mit Nginx',
        '8. Plesk Integration',
        '9. Knowledge Base Administration',
        '10. Support-Levels und Team-Lead',
        '11. Statistiken und Berichte',
        '12. Sicherheit und Datenschutz',
        '13. Backup und Wartung',
        '14. Troubleshooting'
    ]
    for item in toc_items:
        doc.add_paragraph(item, style='List Bullet')

    doc.add_page_break()

    # 1. System Overview
    doc.add_heading('1. Systemübersicht', level=1)
    doc.add_paragraph(
        'Das ML Gruppe Helpdesk System ist eine Django-basierte Anwendung zur Verwaltung von Support-Tickets. '
        'Das System unterstützt mehrere Deployment-Szenarien und Datenbanktypen.'
    )

    doc.add_heading('Benutzerrollen', level=2)
    roles = [
        'Administrator: Vollständiger Zugriff auf alle Funktionen',
        'Support Agent: Kann Tickets verwalten und Knowledge Base Articles erstellen',
        'Team Lead (Level 4): Kann Tickets an andere Agents zuweisen',
        'Customer: Kann nur ihre eigenen Tickets sehen'
    ]
    for role in roles:
        doc.add_paragraph(role, style='List Bullet')

    # 2. User Management
    doc.add_heading('2. Benutzerverwaltung', level=1)

    doc.add_heading('Neue Benutzer erstellen', level=2)
    doc.add_paragraph(
        'Nur Administratoren können neue Benutzer erstellen. Gehen Sie zum Admin-Panel '
        'und navigieren Sie zu "Benutzer" > "Neuer Benutzer".'
    )

    user_fields = [
        'Email: Eindeutige Email-Adresse',
        'Username: Eindeutiger Benutzername',
        'Vorname und Nachname: Vollständiger Name',
        'Rolle: Wählen Sie die entsprechende Rolle',
        'Support Level: Nur für Support Agents (1-4)',
        'Telefonnummer: Kontaktinformation'
    ]
    for field in user_fields:
        doc.add_paragraph(field, style='List Bullet')

    doc.add_heading('Passwortverwaltung', level=2)
    doc.add_paragraph(
        'Neue Kunden erhalten ein initiales Passwort: P@ssw0rd123. '
        'Sie müssen das Passwort bei ihrer ersten Anmeldung ändern.'
    )

    # 3. Database Configuration
    doc.add_heading('3. Datenbank-Konfiguration', level=1)

    doc.add_heading('Datenbank-Optionen', level=2)
    doc.add_paragraph('Das System unterstützt folgende Datenbanken:')

    db_options = [
        'SQLite - Einfach, keine Installation nötig (Standard)',
        'MySQL/MariaDB - Produktionsqualität, weit verbreitet',
        'PostgreSQL - Robust, empfohlen für große Systeme',
        'MongoDB - NoSQL Alternative (eingeschränkte Unterstützung)'
    ]
    for option in db_options:
        doc.add_paragraph(option, style='List Bullet')

    # MySQL/MariaDB
    doc.add_heading('MySQL/MariaDB Konfiguration', level=2)

    doc.add_heading('Installation (Ubuntu/Debian)', level=3)
    doc.add_paragraph('sudo apt-get update')
    doc.add_paragraph('sudo apt-get install mysql-server')
    doc.add_paragraph('sudo mysql_secure_installation')

    doc.add_heading('Datenbank und Benutzer erstellen', level=3)
    doc.add_paragraph('sudo mysql -u root -p')
    doc.add_paragraph('CREATE DATABASE ml_helpdesk CHARACTER SET utf8mb4 COLLATE utf8mb4_unicode_ci;')
    doc.add_paragraph('CREATE USER \'helpdesk\'@\'localhost\' IDENTIFIED BY \'sicheres_passwort\';')
    doc.add_paragraph('GRANT ALL PRIVILEGES ON ml_helpdesk.* TO \'helpdesk\'@\'localhost\';')
    doc.add_paragraph('FLUSH PRIVILEGES;')
    doc.add_paragraph('EXIT;')

    doc.add_heading('Django Konfiguration (settings.py)', level=3)
    db_config = '''
DATABASES = {
    'default': {
        'ENGINE': 'django.db.backends.mysql',
        'NAME': 'ml_helpdesk',
        'USER': 'helpdesk',
        'PASSWORD': 'sicheres_passwort',
        'HOST': 'localhost',
        'PORT': '3306',
        'OPTIONS': {
            'charset': 'utf8mb4',
        }
    }
}
    '''
    doc.add_paragraph(db_config)

    doc.add_heading('Requirements.txt Update', level=3)
    doc.add_paragraph('pip install mysqlclient')

    # PostgreSQL
    doc.add_heading('PostgreSQL Konfiguration', level=2)

    doc.add_heading('Installation (Ubuntu/Debian)', level=3)
    doc.add_paragraph('sudo apt-get update')
    doc.add_paragraph('sudo apt-get install postgresql postgresql-contrib')

    doc.add_heading('Datenbank und Benutzer erstellen', level=3)
    doc.add_paragraph('sudo -u postgres psql')
    doc.add_paragraph('CREATE DATABASE ml_helpdesk;')
    doc.add_paragraph('CREATE USER helpdesk WITH PASSWORD \'sicheres_passwort\';')
    doc.add_paragraph('ALTER ROLE helpdesk SET client_encoding TO \'utf8\';')
    doc.add_paragraph('ALTER ROLE helpdesk SET default_transaction_isolation TO \'read committed\';')
    doc.add_paragraph('ALTER ROLE helpdesk SET default_transaction_deferrable TO on;')
    doc.add_paragraph('GRANT ALL PRIVILEGES ON DATABASE ml_helpdesk TO helpdesk;')
    doc.add_paragraph('\\q')

    doc.add_heading('Django Konfiguration (settings.py)', level=3)
    pg_config = '''
DATABASES = {
    'default': {
        'ENGINE': 'django.db.backends.postgresql',
        'NAME': 'ml_helpdesk',
        'USER': 'helpdesk',
        'PASSWORD': 'sicheres_passwort',
        'HOST': 'localhost',
        'PORT': '5432',
    }
}
    '''
    doc.add_paragraph(pg_config)

    doc.add_heading('Requirements.txt Update', level=3)
    doc.add_paragraph('pip install psycopg2-binary')

    # MongoDB
    doc.add_heading('MongoDB Konfiguration (Optional)', level=2)

    doc.add_heading('Installation (Ubuntu/Debian)', level=3)
    doc.add_paragraph('sudo apt-get update')
    doc.add_paragraph('sudo apt-get install -y mongodb')
    doc.add_paragraph('sudo systemctl start mongodb')

    doc.add_heading('Django Konfiguration', level=3)
    doc.add_paragraph('pip install djongo')
    doc.add_paragraph('pip install pymongo')

    mongo_config = '''
DATABASES = {
    'default': {
        'ENGINE': 'djongo',
        'NAME': 'ml_helpdesk',
        'CLIENT': {
            'host': 'localhost',
            'port': 27017,
        }
    }
}
    '''
    doc.add_paragraph(mongo_config)

    # 4. Caching and Queue
    doc.add_heading('4. Caching und Queue-System', level=1)

    doc.add_heading('Redis Installation', level=2)

    doc.add_heading('Installation (Ubuntu/Debian)', level=3)
    doc.add_paragraph('sudo apt-get update')
    doc.add_paragraph('sudo apt-get install redis-server')
    doc.add_paragraph('sudo systemctl start redis-server')
    doc.add_paragraph('sudo systemctl enable redis-server')

    doc.add_heading('Redis Konfiguration überprüfen', level=3)
    doc.add_paragraph('redis-cli ping')
    doc.add_paragraph('# Sollte PONG zurückgeben')

    doc.add_heading('Django Cache mit Redis', level=3)
    redis_cache = '''
CACHES = {
    'default': {
        'BACKEND': 'django_redis.cache.RedisCache',
        'LOCATION': 'redis://127.0.0.1:6379/1',
        'OPTIONS': {
            'CLIENT_CLASS': 'django_redis.client.DefaultClient',
        }
    }
}

pip install django-redis
    '''
    doc.add_paragraph(redis_cache)

    # Celery
    doc.add_heading('Celery und RabbitMQ', level=2)

    doc.add_heading('RabbitMQ Installation', level=3)
    doc.add_paragraph('sudo apt-get install rabbitmq-server')
    doc.add_paragraph('sudo systemctl start rabbitmq-server')
    doc.add_paragraph('sudo systemctl enable rabbitmq-server')

    doc.add_heading('Celery Installation', level=3)
    doc.add_paragraph('pip install celery')
    doc.add_paragraph('pip install celery[redis]')

    doc.add_heading('Celery Konfiguration (celery.py)', level=3)
    celery_config = '''
import os
from celery import Celery

os.environ.setdefault('DJANGO_SETTINGS_MODULE', 'helpdesk.settings')

app = Celery('helpdesk')
app.config_from_object('django.conf:settings', namespace='CELERY')
app.autodiscover_tasks()

# Broker und Backend konfigurieren
CELERY_BROKER_URL = 'redis://localhost:6379/0'
CELERY_RESULT_BACKEND = 'redis://localhost:6379/0'
CELERY_ACCEPT_CONTENT = ['json']
CELERY_TASK_SERIALIZER = 'json'
    '''
    doc.add_paragraph(celery_config)

    doc.add_heading('Celery Worker starten', level=3)
    doc.add_paragraph('celery -A helpdesk worker -l info')

    doc.add_page_break()

    # 5. Standalone Deployment
    doc.add_heading('5. Standalone Deployment (Apache/Nginx)', level=1)

    doc.add_heading('Apache2 mit Gunicorn', level=2)

    doc.add_heading('Installation', level=3)
    doc.add_paragraph('sudo apt-get install apache2')
    doc.add_paragraph('sudo apt-get install python3-pip python3-venv')
    doc.add_paragraph('pip install gunicorn')

    doc.add_heading('Gunicorn Service erstellen', level=3)
    gunicorn_service = '''
sudo nano /etc/systemd/system/ml-helpdesk.service

[Unit]
Description=ML Helpdesk Gunicorn Application
After=network.target

[Service]
Type=notify
User=www-data
Group=www-data
WorkingDirectory=/var/www/ml-helpdesk
Environment="PATH=/var/www/ml-helpdesk/.venv/bin"
ExecStart=/var/www/ml-helpdesk/.venv/bin/gunicorn \\
    --workers 4 \\
    --worker-class sync \\
    --bind unix:/var/www/ml-helpdesk/gunicorn.sock \\
    helpdesk.wsgi:application

[Install]
WantedBy=multi-user.target

sudo systemctl start ml-helpdesk
sudo systemctl enable ml-helpdesk
    '''
    doc.add_paragraph(gunicorn_service)

    doc.add_heading('Apache VirtualHost konfigurieren', level=3)
    apache_config = '''
sudo nano /etc/apache2/sites-available/ml-helpdesk.conf

<VirtualHost *:80>
    ServerName helpdesk.example.com
    ServerAdmin admin@example.com

    ProxyPreserveHost On
    ProxyPass / unix:/var/www/ml-helpdesk/gunicorn.sock|http://localhost/
    ProxyPassReverse / unix:/var/www/ml-helpdesk/gunicorn.sock|http://localhost/

    <Directory /var/www/ml-helpdesk>
        Require all granted
    </Directory>

    ErrorLog ${APACHE_LOG_DIR}/ml-helpdesk-error.log
    CustomLog ${APACHE_LOG_DIR}/ml-helpdesk-access.log combined
</VirtualHost>

sudo a2enmod proxy
sudo a2enmod proxy_http
sudo a2ensite ml-helpdesk.conf
sudo apache2ctl configtest
sudo systemctl restart apache2
    '''
    doc.add_paragraph(apache_config)

    doc.add_heading('Nginx mit Gunicorn', level=2)

    doc.add_heading('Installation', level=3)
    doc.add_paragraph('sudo apt-get install nginx')
    doc.add_paragraph('sudo systemctl start nginx')
    doc.add_paragraph('sudo systemctl enable nginx')

    doc.add_heading('Nginx konfigurieren', level=3)
    nginx_config = '''
sudo nano /etc/nginx/sites-available/ml-helpdesk

upstream ml_helpdesk {
    server unix:/var/www/ml-helpdesk/gunicorn.sock fail_timeout=0;
}

server {
    listen 80;
    server_name helpdesk.example.com;
    client_max_body_size 100M;

    location / {
        proxy_set_header Host $host;
        proxy_set_header X-Real-IP $remote_addr;
        proxy_set_header X-Forwarded-For $proxy_add_x_forwarded_for;
        proxy_set_header X-Forwarded-Proto $scheme;
        proxy_redirect off;
        proxy_pass http://ml_helpdesk;
    }

    location /static/ {
        alias /var/www/ml-helpdesk/static/;
    }

    location /media/ {
        alias /var/www/ml-helpdesk/media/;
    }

    error_log /var/log/nginx/ml-helpdesk-error.log;
    access_log /var/log/nginx/ml-helpdesk-access.log;
}

sudo ln -s /etc/nginx/sites-available/ml-helpdesk /etc/nginx/sites-enabled/
sudo nginx -t
sudo systemctl restart nginx
    '''
    doc.add_paragraph(nginx_config)

    doc.add_page_break()

    # 6. ISPConfig 3 Apache
    doc.add_heading('6. ISPConfig 3 mit Apache2', level=1)

    doc.add_heading('Vorbedingungen', level=2)
    doc.add_paragraph('ISPConfig 3 muss installiert sein')
    doc.add_paragraph('Apache2 mit mod_proxy und mod_wsgi')
    doc.add_paragraph('Python 3.7+')

    doc.add_heading('Schritt 1: Virtuelle Website in ISPConfig erstellen', level=2)
    doc.add_paragraph('1. ISPConfig Control Panel öffnen (https://server.ip:8080)')
    doc.add_paragraph('2. Login mit Administrator-Anmeldedaten')
    doc.add_paragraph('3. "Sites" > "Websites" > "Website hinzufügen"')
    doc.add_paragraph('4. Domain eingeben (z.B. helpdesk.example.com)')
    doc.add_paragraph('5. Document Root: /var/www/clients/client1/web1/web (Standard)')

    doc.add_heading('Schritt 2: SSH Zugriff aktivieren', level=2)
    doc.add_paragraph('1. In ISPConfig: "System" > "Benutzer"')
    doc.add_paragraph('2. Benutzer auswählen')
    doc.add_paragraph('3. "SSH/Shell login" auf "jailed shell" oder "nologin" setzen')

    doc.add_heading('Schritt 3: Applikation hochladen', level=2)
    doc.add_paragraph('1. FTP/SFTP Zugriff nutzen')
    doc.add_paragraph('2. mini-helpdesk Ordner in Web Root hochladen')
    doc.add_paragraph('3. Berechtigungen: chmod -R 755 /var/www/clients/client1/web1/web/mini-helpdesk')

    doc.add_heading('Schritt 4: Apache VirtualHost anpassen', level=2)
    ispconfig_apache = '''
Editieren Sie folgende Datei:
/etc/apache2/sites-available/[domain].conf

Fügen Sie hinzu:

<VirtualHost *:80>
    ServerName helpdesk.example.com
    DocumentRoot /var/www/clients/client1/web1/web/mini-helpdesk

    <Directory /var/www/clients/client1/web1/web/mini-helpdesk>
        AllowOverride All
        Require all granted
    </Directory>

    ProxyPreserveHost On
    ProxyPass /static/ !
    ProxyPass /media/ !
    ProxyPass / http://127.0.0.1:8001/
    ProxyPassReverse / http://127.0.0.1:8001/
</VirtualHost>
    '''
    doc.add_paragraph(ispconfig_apache)

    doc.add_heading('Schritt 5: Gunicorn als Systemd Service', level=2)
    ispconfig_gunicorn = '''
sudo nano /etc/systemd/system/ml-helpdesk-ispconfig.service

[Unit]
Description=ML Helpdesk Gunicorn
After=network.target

[Service]
Type=notify
User=web1
Group=client1
WorkingDirectory=/var/www/clients/client1/web1/web/mini-helpdesk
Environment="PATH=/var/www/clients/client1/web1/web/mini-helpdesk/.venv/bin"
ExecStart=/var/www/clients/client1/web1/web/mini-helpdesk/.venv/bin/gunicorn \\
    --workers 2 \\
    --bind 127.0.0.1:8001 \\
    helpdesk.wsgi:application

[Install]
WantedBy=multi-user.target

sudo systemctl daemon-reload
sudo systemctl start ml-helpdesk-ispconfig
sudo systemctl enable ml-helpdesk-ispconfig
    '''
    doc.add_paragraph(ispconfig_gunicorn)

    doc.add_page_break()

    # 7. ISPConfig 3 Nginx
    doc.add_heading('7. ISPConfig 3 mit Nginx', level=1)

    doc.add_heading('Vorbedingungen', level=2)
    doc.add_paragraph('ISPConfig 3 mit Nginx')
    doc.add_paragraph('Python 3.7+')
    doc.add_paragraph('Gunicorn installiert')

    doc.add_heading('Schritt 1: Virtuelle Website erstellen', level=2)
    doc.add_paragraph('1. ISPConfig öffnen')
    doc.add_paragraph('2. "Sites" > "Websites" > "Website hinzufügen"')
    doc.add_paragraph('3. Unter "Web Server": Nginx auswählen')
    doc.add_paragraph('4. Domain und Document Root konfigurieren')

    doc.add_heading('Schritt 2: Nginx Server Block bearbeiten', level=2)
    ispconfig_nginx = '''
sudo nano /etc/nginx/sites-available/[domain].conf

Beispiel-Konfiguration:

upstream ml_helpdesk {
    server 127.0.0.1:8002;
}

server {
    listen 80;
    server_name helpdesk.example.com;
    client_max_body_size 100M;
    root /var/www/clients/client1/web3/web/mini-helpdesk;

    location /static/ {
        alias /var/www/clients/client1/web3/web/mini-helpdesk/static/;
        expires 30d;
    }

    location /media/ {
        alias /var/www/clients/client1/web3/web/mini-helpdesk/media/;
        expires 7d;
    }

    location / {
        proxy_pass http://ml_helpdesk;
        proxy_set_header Host $host;
        proxy_set_header X-Real-IP $remote_addr;
        proxy_set_header X-Forwarded-For $proxy_add_x_forwarded_for;
        proxy_set_header X-Forwarded-Proto $scheme;
    }
}

sudo systemctl reload nginx
    '''
    doc.add_paragraph(ispconfig_nginx)

    doc.add_heading('Schritt 3: Gunicorn für ISPConfig Nginx', level=2)
    ispconfig_nginx_gunicorn = '''
sudo nano /etc/systemd/system/ml-helpdesk-nginx.service

[Unit]
Description=ML Helpdesk Gunicorn for Nginx
After=network.target

[Service]
Type=notify
User=web3
Group=client1
WorkingDirectory=/var/www/clients/client1/web3/web/mini-helpdesk
Environment="PATH=/var/www/clients/client1/web3/web/mini-helpdesk/.venv/bin"
ExecStart=/var/www/clients/client1/web3/web/mini-helpdesk/.venv/bin/gunicorn \\
    --workers 2 \\
    --bind 127.0.0.1:8002 \\
    helpdesk.wsgi:application

[Install]
WantedBy=multi-user.target

sudo systemctl daemon-reload
sudo systemctl start ml-helpdesk-nginx
sudo systemctl enable ml-helpdesk-nginx
    '''
    doc.add_paragraph(ispconfig_nginx_gunicorn)

    doc.add_page_break()

    # 8. Plesk
    doc.add_heading('8. Plesk Integration', level=1)

    doc.add_heading('Vorbedingungen', level=2)
    doc.add_paragraph('Plesk 17.8+ installiert')
    doc.add_paragraph('Python 3.7+ aktiviert')
    doc.add_paragraph('Node.js Modul installiert (optional)')

    doc.add_heading('Schritt 1: Subdomain/Domain erstellen', level=2)
    doc.add_paragraph('1. Plesk Control Panel öffnen')
    doc.add_paragraph('2. "Domains" > Domain auswählen')
    doc.add_paragraph('3. "Subdomains" > "Add Subdomain"')
    doc.add_paragraph('4. Subdomain Name eingeben (z.B. "helpdesk")')
    doc.add_paragraph('5. Bestätigen')

    doc.add_heading('Schritt 2: Python Anwendung hochladen', level=2)
    plesk_upload = '''
1. FTP/SFTP nutzen oder Files Manager in Plesk
2. mini-helpdesk Ordner hochladen in: /httpdocs/helpdesk/
3. Dateiberechtigungen setzen: 755 für Ordner, 644 für Dateien
    '''
    doc.add_paragraph(plesk_upload)

    doc.add_heading('Schritt 3: Virtual Environment erstellen', level=2)
    plesk_venv = '''
SSH als root:

cd /var/www/vhosts/example.com/httpdocs/helpdesk
python3 -m venv .venv
source .venv/bin/activate
pip install -r requirements.txt
pip install gunicorn
    '''
    doc.add_paragraph(plesk_venv)

    doc.add_heading('Schritt 4: Node.js Anwendung in Plesk konfigurieren', level=2)
    plesk_nodejs = '''
1. In Plesk: "Domains" > Domain > "Node.js"
2. "Add Node.js Application" klicken
3. Einstellungen:
   - Path: /httpdocs/helpdesk/
   - Document root: /httpdocs/helpdesk/public
   - Application mode: production
   - Port: 3000 (oder beliebiger Port)
   - NPM Scripts: verwende gunicorn statt npm

4. Alternativ: Custom Startup Script
   Script: /var/www/vhosts/example.com/.venv/bin/gunicorn \\
           --workers 2 --bind 127.0.0.1:3000 helpdesk.wsgi:application
    '''
    doc.add_paragraph(plesk_nodejs)

    doc.add_heading('Schritt 5: Apache/Nginx Proxy konfigurieren', level=2)
    plesk_proxy = '''
Plesk erstellt automatisch Proxy-Regeln. Falls nicht:

Apache (httpd.conf oder htaccess):
RewriteEngine On
RewriteRule ^(.*)$ http://127.0.0.1:3000/$1 [P,L]

Nginx (wird in Plesk-Interface konfiguriert)
    '''
    doc.add_paragraph(plesk_proxy)

    doc.add_page_break()

    # 9. Knowledge Base
    doc.add_heading('9. Knowledge Base Administration', level=1)

    doc.add_heading('Artikel erstellen', level=2)
    doc.add_paragraph(
        'Support Agents und Administratoren können Knowledge Base Artikel erstellen. '
        'Artikel sollten häufige Probleme und deren Lösungen dokumentieren.'
    )

    kb_fields = [
        'Titel: Aussagekräftiger Titel',
        'Kategorie: Thematische Zuordnung',
        'Inhalt: Detaillierte Anleitung (mit RichText-Editor)',
        'Tags: Für bessere Auffindbarkeit',
        'Veröffentlicht: Status des Artikels'
    ]
    for field in kb_fields:
        doc.add_paragraph(field, style='List Bullet')

    # 10. Support Levels
    doc.add_heading('10. Support-Levels und Team-Lead', level=1)

    doc.add_heading('Support Levels', level=2)
    doc.add_paragraph('Es gibt 4 Support Levels für Agents:')
    levels = [
        'Level 1: Junior Support (kann nur einfache Tickets lösen)',
        'Level 2: Senior Support (normale Tickets)',
        'Level 3: Expert Support (komplexe Probleme)',
        'Level 4: Team Lead (kann Tickets zuweisen und delegieren)'
    ]
    for level in levels:
        doc.add_paragraph(level, style='List Bullet')

    # 11. Statistics
    doc.add_heading('11. Statistiken und Berichte', level=1)
    doc.add_paragraph(
        'Das System bietet umfangreiche Statistiken zur Überwachung der Support-Qualität. '
        'Nur Administratoren und Team Leads haben Zugriff auf Statistiken.'
    )

    # 12. Security
    doc.add_heading('12. Sicherheit und Datenschutz', level=1)

    doc.add_heading('SSL/TLS Zertifikat', level=2)
    doc.add_paragraph('Verwenden Sie Let\'s Encrypt für kostenlose Zertifikate:')
    doc.add_paragraph('sudo apt-get install certbot python3-certbot-apache')
    doc.add_paragraph('sudo certbot certonly --apache -d helpdesk.example.com')

    doc.add_heading('Firewall Konfiguration', level=2)
    doc.add_paragraph('sudo ufw allow 22/tcp')
    doc.add_paragraph('sudo ufw allow 80/tcp')
    doc.add_paragraph('sudo ufw allow 443/tcp')
    doc.add_paragraph('sudo ufw enable')

    doc.add_heading('Django Security Settings', level=2)
    security_settings = '''
# settings.py

DEBUG = False
ALLOWED_HOSTS = ['helpdesk.example.com']

# HTTPS
SECURE_SSL_REDIRECT = True
SESSION_COOKIE_SECURE = True
CSRF_COOKIE_SECURE = True
SECURE_BROWSER_XSS_FILTER = True
SECURE_CONTENT_SECURITY_POLICY = {
    'default-src': ("'self'",),
}
    '''
    doc.add_paragraph(security_settings)

    # 13. Backup
    doc.add_heading('13. Backup und Wartung', level=1)

    doc.add_heading('Automatische Datenbank-Backups', level=2)
    doc.add_paragraph('#!/bin/bash')
    doc.add_paragraph('BACKUP_DIR="/backups/ml-helpdesk"')
    doc.add_paragraph('DATE=$(date +%Y%m%d_%H%M%S)')
    doc.add_paragraph('# PostgreSQL Backup')
    doc.add_paragraph('pg_dump ml_helpdesk > $BACKUP_DIR/db_backup_$DATE.sql')
    doc.add_paragraph('gzip $BACKUP_DIR/db_backup_$DATE.sql')
    doc.add_paragraph('# MySQL Backup')
    doc.add_paragraph('mysqldump -u helpdesk -p ml_helpdesk > $BACKUP_DIR/db_backup_$DATE.sql')

    doc.add_heading('Cron Job für tägliche Backups', level=2)
    doc.add_paragraph('crontab -e')
    doc.add_paragraph('0 2 * * * /usr/local/bin/backup-helpdesk.sh')

    # 14. Troubleshooting
    doc.add_heading('14. Troubleshooting', level=1)

    doc.add_heading('Common Issues', level=2)

    issues = [
        ('502 Bad Gateway Error', 'Gunicorn nicht erreichbar. Überprüfen Sie: systemctl status ml-helpdesk'),
        ('Datenbankfehler', 'Überprüfen Sie Datenbank-Verbindung und Credentials in settings.py'),
        ('Static Files nicht geladen', 'Führen Sie aus: python manage.py collectstatic'),
        ('Email funktioniert nicht', 'Überprüfen Sie SMTP-Einstellungen und Firewall-Regeln'),
        ('Permission denied Fehler', 'Überprüfen Sie Dateiberechtigungen: chmod -R 755 /var/www/ml-helpdesk')
    ]

    for issue, solution in issues:
        doc.add_heading(issue, level=3)
        doc.add_paragraph(solution)

    # Footer
    doc.add_paragraph()
    footer_para = doc.add_paragraph()
    footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer_run = footer_para.add_run('© 2025 ML Gruppe Helpdesk System - Erweiterte Administrationshandbuch')
    footer_run.font.size = Pt(8)
    footer_run.font.italic = True

    doc.save('Administrationshandbuch_ML_Helpdesk_ERWEITERT.docx')
    return 'Administrationshandbuch_ML_Helpdesk_ERWEITERT.docx'

if __name__ == '__main__':
    print("[ERSTELLE] Erweitertes Administrationshandbuch...\n")
    file = create_extended_admin_manual()
    print(f"[OK] {file} erstellt")
    print(f"\n[SUCCESS] Erweitertes Administrationshandbuch fertig!")
    print(f"\nDatei:")
    print(f"  - Administrationshandbuch_ML_Helpdesk_ERWEITERT.docx")
